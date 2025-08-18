#!/usr/bin/env python3
"""
Inspect the generated document to find duplications
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from logic.srs_generator import generate_srs_docx
from docx import Document
import tempfile

def inspect_document_structure():
    """Generate a document and inspect its structure"""
    print("🔍 Inspecting Document Structure")
    print("=" * 60)
    
    # Create test headings
    test_headings = [
        {
            'heading': 'System Overview',
            'purpose': 'Provide system overview',
            'category': 'Introduction',
            'source': 'Standard Template',
            'userPrompt': ''
        },
        {
            'heading': 'Database Schema',
            'purpose': 'Show database structure',
            'category': 'Custom',
            'source': 'Custom Section',
            'userPrompt': 'Generate a simple ER diagram'
        }
    ]
    
    print(f"📋 Input headings:")
    for i, heading in enumerate(test_headings, 1):
        print(f"  {i}. {heading['heading']} ({heading['category']})")
    
    # Generate document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        output_path = temp_file.name
    
    try:
        print(f"\n🚀 Generating document...")
        generate_srs_docx(test_headings, output_path, "Test meeting context")
        
        if os.path.exists(output_path):
            print(f"✅ Document generated: {output_path}")
            
            # Open and inspect the document
            print(f"\n📄 Inspecting document structure...")
            doc = Document(output_path)
            
            print(f"\n📊 Document Analysis:")
            print(f"   Total paragraphs: {len(doc.paragraphs)}")
            
            # Find all headings
            headings_found = []
            toc_entries = []
            content_headings = []
            
            in_toc_section = False
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                style = para.style.name if para.style else "No Style"
                
                # Track TOC section
                if text == "Table of Contents":
                    in_toc_section = True
                    continue
                elif text == "Software Requirements Specification":
                    in_toc_section = False
                    continue
                
                # Collect headings
                if style in ['SRS Heading 1', 'SRS Heading 2', 'Heading 1', 'Heading 2']:
                    headings_found.append({
                        'line': i + 1,
                        'text': text,
                        'style': style,
                        'in_toc': in_toc_section
                    })
                    
                    if in_toc_section:
                        toc_entries.append(text)
                    else:
                        content_headings.append(text)
                
                # Also check for numbered entries that might be TOC items
                if in_toc_section and (text.startswith('1.') or text.startswith('2.') or text.startswith('3.')):
                    toc_entries.append(text)
            
            print(f"\n📋 All Headings Found ({len(headings_found)}):")
            for heading in headings_found:
                location = "TOC" if heading['in_toc'] else "Content"
                print(f"   Line {heading['line']:3d}: [{location:7s}] {heading['style']:15s} | {heading['text']}")
            
            print(f"\n📑 TOC Entries ({len(toc_entries)}):")
            for i, entry in enumerate(toc_entries, 1):
                print(f"   {i}. {entry}")
            
            print(f"\n📄 Content Headings ({len(content_headings)}):")
            for i, heading in enumerate(content_headings, 1):
                print(f"   {i}. {heading}")
            
            # Check for duplicates
            print(f"\n🔍 Duplicate Analysis:")
            
            # Check TOC duplicates
            toc_counts = {}
            for entry in toc_entries:
                clean_entry = entry.replace('1. ', '').replace('2. ', '').replace('3. ', '').strip()
                if clean_entry:
                    toc_counts[clean_entry] = toc_counts.get(clean_entry, 0) + 1
            
            toc_duplicates = {k: v for k, v in toc_counts.items() if v > 1}
            
            # Check content duplicates
            content_counts = {}
            for heading in content_headings:
                if heading:
                    content_counts[heading] = content_counts.get(heading, 0) + 1
            
            content_duplicates = {k: v for k, v in content_counts.items() if v > 1}
            
            if toc_duplicates:
                print(f"   ❌ TOC Duplicates found:")
                for item, count in toc_duplicates.items():
                    print(f"      '{item}' appears {count} times")
            else:
                print(f"   ✅ No TOC duplicates found")
            
            if content_duplicates:
                print(f"   ❌ Content Duplicates found:")
                for item, count in content_duplicates.items():
                    print(f"      '{item}' appears {count} times")
            else:
                print(f"   ✅ No content duplicates found")
            
            return len(toc_duplicates) == 0 and len(content_duplicates) == 0
            
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
    
    finally:
        # Keep file for inspection
        if os.path.exists(output_path):
            print(f"\n📁 Document saved for manual inspection: {output_path}")

if __name__ == "__main__":
    print("🚀 Document Structure Inspector")
    print("=" * 70)
    
    success = inspect_document_structure()
    
    print("\n" + "=" * 70)
    if success:
        print("✅ No duplicates found in document structure!")
    else:
        print("❌ Duplicates detected in document structure!")
        print("\n💡 This helps identify where the duplication is happening:")
        print("   • TOC duplicates = Issue in add_table_of_contents function")
        print("   • Content duplicates = Issue in add_content_sections function")
    print("=" * 70)
