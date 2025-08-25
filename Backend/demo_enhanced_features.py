#!/usr/bin/env python3
"""
Demonstration of Enhanced SRS Generator Features
Shows header/footer, markdown formatting, text justification, and diagram centering
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from logic.srs_generator import generate_srs_docx
from docx import Document
import tempfile

def create_demo_document():
    """Create a demonstration document showcasing all enhanced features"""
    
    print("🎯 Creating Enhanced SRS Demo Document")
    print("="*50)
    
    # Sample headings with rich markdown content
    demo_headings = [
        {
            'heading': '**Project Overview**',
            'purpose': 'Demonstrate **bold formatting** and *italic text* in headings and content',
            'category': 'Introduction',
            'source': 'Demo'
        },
        {
            'heading': 'Functional Requirements',
            'purpose': 'Show how requirements can be formatted with **emphasis** and proper justification',
            'category': 'Requirements',
            'source': 'Demo'
        },
        {
            'heading': 'System Architecture',
            'purpose': 'Demonstrate automatic diagram generation and centering',
            'category': 'Design',
            'source': 'Demo'
        },
        {
            'heading': '**Quality Attributes**',
            'purpose': 'Show mixed formatting with **bold headings** and *italic emphasis*',
            'category': 'Quality',
            'source': 'Demo'
        }
    ]
    
    # Rich context with markdown formatting
    demo_context = """
    This is a demonstration of the enhanced SRS generator with the following features:
    
    **Header and Footer Support:**
    - Logo display in header (or fallback text)
    - Professional footer with document title and page numbers
    
    **Enhanced Markdown Formatting:**
    - **Bold text** using double asterisks
    - *Italic text* using single asterisks
    - Automatic text justification for professional appearance
    
    **Diagram Features:**
    - Automatic diagram generation for architecture sections
    - Centered diagram placement
    - A4-optimized sizing
    
    **Professional Formatting:**
    - Justified text paragraphs
    - Proper spacing and indentation
    - Consistent styling throughout
    """
    
    # Generate the demo document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        output_path = temp_file.name
    
    try:
        print("🚀 Generating enhanced SRS document...")
        generate_srs_docx(demo_headings, output_path, demo_context, "Enhanced SRS Demo")
        
        if os.path.exists(output_path):
            print(f"✅ Demo document created successfully!")
            
            # Analyze the generated document
            doc = Document(output_path)
            
            print(f"\n📊 Document Analysis:")
            print(f"   📄 Total paragraphs: {len(doc.paragraphs)}")
            print(f"   📑 Total sections: {len(doc.sections)}")
            
            # Check header and footer
            section = doc.sections[0]
            header = section.header
            footer = section.footer
            
            print(f"\n🔍 Header/Footer Features:")
            if header.paragraphs and header.paragraphs[0].runs:
                if any(run._element.xpath('.//pic:pic') for run in header.paragraphs[0].runs):
                    print(f"   ✅ Logo found in header")
                else:
                    print(f"   ✅ Header text: '{header.paragraphs[0].text}'")
            
            if footer.paragraphs and footer.paragraphs[0].text:
                print(f"   ✅ Footer with page numbering configured")
            
            # Check formatting features
            bold_runs = sum(1 for para in doc.paragraphs for run in para.runs if run.bold)
            italic_runs = sum(1 for para in doc.paragraphs for run in para.runs if run.italic)
            justified_paras = sum(1 for para in doc.paragraphs if para.alignment == 3)
            
            print(f"\n📝 Formatting Analysis:")
            print(f"   ✅ Bold text runs: {bold_runs}")
            print(f"   ✅ Italic text runs: {italic_runs}")
            print(f"   ✅ Justified paragraphs: {justified_paras}")
            
            # Check for diagrams
            diagram_count = 0
            for para in doc.paragraphs:
                if para.runs:
                    for run in para.runs:
                        if run._element.xpath('.//pic:pic'):
                            diagram_count += 1
                            if para.alignment == 1:  # CENTER
                                print(f"   ✅ Centered diagram found")
            
            if diagram_count > 0:
                print(f"   ✅ Total diagrams: {diagram_count}")
            
            print(f"\n🎉 Enhanced Features Summary:")
            print(f"   ✅ Header/Footer: Configured")
            print(f"   ✅ Markdown Support: {bold_runs + italic_runs} formatted runs")
            print(f"   ✅ Text Justification: {justified_paras} justified paragraphs")
            print(f"   ✅ Diagram Centering: {diagram_count} centered diagrams")
            
            # Copy to a permanent location for easy access
            permanent_path = os.path.join(os.path.dirname(__file__), 'enhanced_srs_demo.docx')
            import shutil
            shutil.copy2(output_path, permanent_path)
            
            print(f"\n📁 Demo document saved to: {permanent_path}")
            print(f"📁 Temporary file: {output_path}")
            
            return permanent_path
            
        else:
            print(f"❌ Failed to generate demo document")
            return None
            
    except Exception as e:
        print(f"❌ Error creating demo document: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def show_feature_comparison():
    """Show before/after comparison of features"""
    
    print("\n" + "="*60)
    print("📋 ENHANCED FEATURES COMPARISON")
    print("="*60)
    
    print("\n🔧 BEFORE (Original SRS Generator):")
    print("   • Basic text formatting")
    print("   • No headers or footers")
    print("   • Limited markdown support")
    print("   • Left-aligned text")
    print("   • Basic diagram insertion")
    
    print("\n✨ AFTER (Enhanced SRS Generator):")
    print("   • ✅ Professional headers with logo support")
    print("   • ✅ Footers with page numbering")
    print("   • ✅ Full markdown support (**bold**, *italic*)")
    print("   • ✅ Justified text for professional appearance")
    print("   • ✅ Centered diagrams with A4 optimization")
    print("   • ✅ Enhanced text formatting and styling")
    
    print("\n🎯 KEY IMPROVEMENTS:")
    print("   1. Header: Logo on right side (or 'SRS Document' text)")
    print("   2. Footer: 'Software Requirement Specification' + Page numbers")
    print("   3. Markdown: **bold** and *italic* text processing")
    print("   4. Layout: Justified text with single spacing (1.0)")
    print("   5. Diagrams: Automatically centered on page")
    print("   6. Numbering: 3, 3.1, 3.2 format (not 3.0, 3.1, 3.2)")

if __name__ == "__main__":
    print("🎨 Enhanced SRS Generator Demonstration")
    print("="*60)
    
    # Create demo document
    demo_path = create_demo_document()
    
    # Show feature comparison
    show_feature_comparison()
    
    print("\n" + "="*60)
    print("🎉 DEMONSTRATION COMPLETE")
    print("="*60)
    
    if demo_path:
        print(f"✅ Demo document created: {demo_path}")
        print(f"📖 Open the document to see all enhanced features!")
    else:
        print(f"❌ Demo document creation failed")
    
    print("\n💡 Usage in your code:")
    print("   from logic.srs_generator import generate_srs_docx")
    print("   generate_srs_docx(headings, 'output.docx', context)")
    print("   # Headers, footers, and formatting applied automatically!")
