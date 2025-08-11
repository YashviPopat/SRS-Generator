import PyPDF2
import docx
import re
from typing import List, Dict, Any
import io
import google.generativeai as genai
import os

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyAv7B9Y6kivXsOkKUgcrJF0Z3Z3wVkcuFA")

class DocumentExtractor:
    """Extract headings and content from various document types"""
    
    @staticmethod
    def extract_from_pdf(pdf_content: bytes) -> Dict[str, Any]:
        """
        Extract headings and content from PDF document
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Dictionary containing extracted headings and content
        """
        try:
            # Create PDF reader from bytes
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            extracted_data = {
                'headings': [],
                'content': '',
                'pages': len(pdf_reader.pages),
                'extraction_method': 'pdf'
            }
            
            # Extract text from all pages
            full_text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                full_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            extracted_data['content'] = full_text
            
            # Extract headings using regex patterns
            headings = DocumentExtractor._extract_headings_from_text(full_text)
            extracted_data['headings'] = headings
            
            return extracted_data
            
        except Exception as e:
            raise Exception(f"Failed to extract from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(docx_content: bytes) -> Dict[str, Any]:
        """
        Extract headings and content from Word document
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Dictionary containing extracted headings and content
        """
        try:
            # Create document from bytes
            doc_file = io.BytesIO(docx_content)
            doc = docx.Document(doc_file)
            
            extracted_data = {
                'headings': [],
                'content': '',
                'paragraphs': len(doc.paragraphs),
                'extraction_method': 'docx'
            }
            
            # Extract headings from document structure
            headings = []
            full_text = ""
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text += text + "\n"
                    
                    # Check if paragraph is a heading based on style
                    if paragraph.style.name.startswith('Heading'):
                        heading_level = int(paragraph.style.name.replace('Heading ', ''))
                        headings.append({
                            'text': text,
                            'level': heading_level,
                            'style': paragraph.style.name
                        })
            
            # Also extract headings using regex patterns
            regex_headings = DocumentExtractor._extract_headings_from_text(full_text)
            
            # Combine both methods
            all_headings = headings + regex_headings
            extracted_data['headings'] = all_headings
            extracted_data['content'] = full_text
            
            return extracted_data
            
        except Exception as e:
            raise Exception(f"Failed to extract from DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_text(text_content: bytes) -> Dict[str, Any]:
        """
        Extract headings from plain text document
        
        Args:
            text_content: Text file content as bytes
            
        Returns:
            Dictionary containing extracted headings and content
        """
        try:
            text = text_content.decode('utf-8', errors='ignore')
            
            extracted_data = {
                'headings': [],
                'content': text,
                'extraction_method': 'text'
            }
            
            # Extract headings using regex patterns
            headings = DocumentExtractor._extract_headings_from_text(text)
            extracted_data['headings'] = headings
            
            return extracted_data
            
        except Exception as e:
            raise Exception(f"Failed to extract from text: {str(e)}")
    
    @staticmethod
    def _extract_headings_from_text(text: str) -> List[Dict[str, Any]]:
        """
        Extract headings from text using regex patterns
        
        Args:
            text: Text content to analyze
            
        Returns:
            List of extracted headings with their levels and positions
        """
        headings = []
        
        # Split text into lines and process each line
        lines = text.split('\n')
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Pattern 1: Numbered headings like "1. Introduction" or "1.1. Document Purpose"
            numbered_match = re.match(r'^(\d+\.(?:\d+)*)\s+(.+?)(?:\s*\.{2,}\s*\d+)?$', line)
            if numbered_match:
                number = numbered_match.group(1)
                heading_text = numbered_match.group(2).strip()
                
                # Determine level based on number of dots
                level = number.count('.')
                
                if heading_text and len(heading_text) > 3:
                    headings.append({
                        'text': heading_text,
                        'level': level,
                        'line_number': line_num + 1,
                        'pattern': 'numbered'
                    })
                continue
            
            # Pattern 2: Section headings like "Introduction" or "System Architecture"
            section_match = re.match(r'^([A-Z][A-Za-z\s&]+?)(?:\s*\.{2,}\s*\d+)?$', line)
            if section_match:
                heading_text = section_match.group(1).strip()
                
                # Skip if it's just a page number or too short
                if (heading_text and len(heading_text) > 3 and 
                    not re.match(r'^Page \d+', heading_text) and
                    not re.match(r'^\d+$', heading_text)):
                    
                    # Determine level based on content
                    level = 1
                    if any(word in heading_text.lower() for word in ['purpose', 'objective', 'stakeholder']):
                        level = 2
                    elif any(word in heading_text.lower() for word in ['component', 'service', 'feature']):
                        level = 3
                    
                    headings.append({
                        'text': heading_text,
                        'level': level,
                        'line_number': line_num + 1,
                        'pattern': 'section'
                    })
                continue
            
            # Pattern 3: Subsection headings like "Document Purpose" or "Project Objectives"
            subsection_match = re.match(r'^([A-Z][A-Za-z\s&]+?)(?:\s*[:.]?\s*(.+))?$', line)
            if subsection_match:
                heading_text = subsection_match.group(1).strip()
                
                # Skip if it's just a page number or too short
                if (heading_text and len(heading_text) > 3 and 
                    not re.match(r'^Page \d+', heading_text) and
                    not re.match(r'^\d+$', heading_text) and
                    len(heading_text) < 100):  # Avoid very long lines
                    
                    level = 2
                    headings.append({
                        'text': heading_text,
                        'level': level,
                        'line_number': line_num + 1,
                        'pattern': 'subsection'
                    })
                continue
        
        # Remove duplicates while preserving order
        seen = set()
        unique_headings = []
        for heading in headings:
            # Clean up heading text
            clean_text = re.sub(r'\s+', ' ', heading['text']).strip()
            if clean_text.lower() not in seen and len(clean_text) > 3:
                seen.add(clean_text.lower())
                heading['text'] = clean_text
                unique_headings.append(heading)
        
        return unique_headings
    
    @staticmethod
    def extract_document_headings(file_content: bytes, file_type: str) -> Dict[str, Any]:
        """
        Main method to extract headings from any supported document type
        
        Args:
            file_content: File content as bytes
            file_type: Type of document ('pdf', 'docx', 'txt', etc.)
            
        Returns:
            Dictionary containing extracted data
        """
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return DocumentExtractor.extract_from_pdf(file_content)
        elif file_type in ['docx', 'doc']:
            return DocumentExtractor.extract_from_docx(file_content)
        elif file_type in ['txt', 'md']:
            return DocumentExtractor.extract_from_text(file_content)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def format_headings_for_api(headings: List[Dict[str, Any]], source_file: str) -> List[Dict[str, str]]:
        """
        Format extracted headings for API response
        
        Args:
            headings: List of extracted headings
            source_file: Name of the source file
            
        Returns:
            List of formatted headings
        """
        formatted_headings = []
        
        for heading in headings:
            # Generate a purpose based on the heading text
            purpose = DocumentExtractor._generate_heading_purpose(heading['text'])
            
            formatted_headings.append({
                'heading': heading['text'],
                'purpose': purpose,
                'level': heading.get('level', 1),
                'source': source_file
            })
        
        return formatted_headings
    
    @staticmethod
    def _generate_heading_purpose(heading_text: str) -> str:
        """
        Generate a purpose description for a heading
        
        Args:
            heading_text: The heading text
            
        Returns:
            Generated purpose description
        """
        heading_lower = heading_text.lower()
        
        # Common purpose mappings
        purpose_mappings = {
            'introduction': 'Briefly describe the system and its main objectives',
            'overview': 'Provide a high-level overview of the project',
            'requirements': 'Define the functional and non-functional requirements',
            'specifications': 'Detailed technical specifications and constraints',
            'design': 'System architecture and design decisions',
            'implementation': 'Implementation details and coding standards',
            'testing': 'Testing strategy and test cases',
            'deployment': 'Deployment procedures and configuration',
            'conclusion': 'Summary and final remarks',
            'references': 'List of references and citations',
            'appendix': 'Additional supporting information',
            'architecture': 'System architecture and component design',
            'database': 'Database design and data models',
            'api': 'API specifications and endpoints',
            'security': 'Security requirements and protocols',
            'user interface': 'UI/UX specifications and wireframes',
            'performance': 'Performance requirements and benchmarks',
            'scalability': 'Scalability considerations and strategies'
        }
        
        # Check for exact matches
        for key, purpose in purpose_mappings.items():
            if key in heading_lower:
                return purpose
        
        # Check for partial matches
        for key, purpose in purpose_mappings.items():
            if any(word in heading_lower for word in key.split()):
                return purpose
        
        # Default purpose
        return f"Describe the {heading_text.lower()} aspects of the system" 

def get_gemini_srs_headings_from_transcript(transcript_text: str) -> dict:
    """
    Call Gemini 2.5 Flash model to generate SRS headings and subheadings (nested, with purposes) from a meeting transcript.
    Returns a nested dict: {heading: {subheading: purpose, ...}, ...}
    """
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel("gemini-2.5-flash")
    
    prompt = f"""
    You are an expert software requirements analyst. Analyze the following meeting transcript and generate SRS headings that are SPECIFICALLY relevant to the content discussed in the meeting, Only include which are actually written in SRS document, not include like in meeting discussion or as meeting transcript because that type of sentences are not mentioned in SRS document.

    MEETING TRANSCRIPT:
    {transcript_text}

    CRITICAL INSTRUCTIONS:
    1. ONLY generate headings that are directly related to topics discussed in the meeting
    2. Use specific project names, technologies, and requirements mentioned in the transcript
    3. Create headings that reflect the actual meeting content, not generic SRS templates
    4. If the meeting discusses specific features, systems, or requirements, create headings for those
    5. Base your suggestions on the actual meeting discussions, decisions, and requirements mentioned

    EXAMPLES:
    - If the meeting discusses "user authentication system", create a heading like "User Authentication System"
    - If the meeting mentions "database requirements", create a heading like "Database Requirements"
    - If the meeting discusses "API integration", create a heading like "API Integration Requirements"

    Return the result as a JSON object in this format: {{"Heading": {{"Subheading": "Purpose", ...}}, ...}}
    If a heading has no subheadings, use a string as the value for its purpose.

    IMPORTANT: Only include headings that are actually mentioned or implied in the meeting transcript above.
    """
    
    try:
        response = model.generate_content(prompt)
        import json
        import re
        
        # Extract JSON from the response
        match = re.search(r'\{[\s\S]*\}', response.text)
        if match:
            try:
                result = json.loads(match.group(0))
                print(f"✅ Gemini generated {len(result)} heading categories")
                return result
            except Exception as e:
                print(f"❌ Failed to parse Gemini JSON: {e}")
                pass
        
        # Fallback: try to parse the whole response
        try:
            result = json.loads(response.text)
            print(f"✅ Gemini generated {len(result)} heading categories (fallback)")
            return result
        except Exception as e:
            print(f"❌ Failed to parse Gemini response: {e}")
            return {}
            
    except Exception as e:
        print(f"❌ Gemini API error: {e}")
        return {} 