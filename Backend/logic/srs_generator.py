#!/usr/bin/env python3
"""
SRS Document Generator
Generates DOCX documents from selected headings and their content
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import os
import subprocess
import tempfile
import json
from datetime import datetime
from typing import List, Dict, Any
import google.generativeai as genai
import re
import hashlib
import markdown
from bs4 import BeautifulSoup

# Global variable to store generated diagrams for the current document
current_document_diagrams = []

# Global counter to limit total diagrams per document (maximum 3)
total_diagrams_generated = 0

# Global set to track diagram signatures and prevent duplicates
generated_diagram_signatures = set()
# Global list to track tokenized signatures for similarity detection
generated_diagram_token_signatures = []

# Global set to track diagram signatures and prevent duplicates
generated_diagram_signatures = set()

# Specific prompts for each standard heading
HEADING_SPECIFIC_PROMPTS = {
    "Document Purpose": """
    Give me the Document Purpose in 1 paragraph of 4 to 5 lines.
    """,

    "Project Purpose": """
    Give me the Project Purpose in 1 paragraph of 4 to 5 lines.
    """,

    "Project Objectives": """
    Give me Project Objectives as point & sub-point wise with one-line description each.
    """,

    "Project Stakeholders": """
    List Project Stakeholders in a table with Role, Name, and Email.
    """,

    "Scope": """
    Give me table with columns Module, Sub Module, Features columns for each usecases.
    """,

    "Assumptions": """
    Give me Assumptions as a numbered list with each assumption in 1–2 lines.
    """,
    
    "Acronyms and Abbreviations" : """Give me Acronyms and Abbreviations in a two-column table: Acronym and Full Form.""",
    
    "User Personas" : """Give me each User Persona in 1 paragraph of 5–6 lines covering role, responsibilities, challenges, and benefits.""",

    "Functional Requirements": """
    Specify detailed functional requirements with preconditions, postconditions, normal flows, alternate flows, exception handling scenarios for each use case. Include field validations, business rules, role-based permissions, audit requirements, and data processing rules for each uses.
    """,

    "Non-Functional Requirements": """
    Define quantified, testable non-functional requirements including performance metrics, scalability targets, reliability standards, security requirements, usability criteria, maintainability standards, and compliance requirements.
    """,

    "System Architecture": """
    Describe the system architecture with component interactions, data flows, integration patterns, deployment topology, and technology stack. Include architectural patterns, design principles, scalability considerations, security architecture, and operational architecture.
    """,

    "Hardware Requirements": """
    Specify detailed hardware requirements including server specifications, network infrastructure, storage requirements, client device specifications, and environmental requirements. Include capacity planning, redundancy requirements, and hardware lifecycle considerations.
    """,

    "Software Requirements": """
    Define software requirements including operating systems, middleware, databases, development tools, runtime environments, and third-party software dependencies. Include version requirements, licensing considerations, and compatibility matrices.
    """,

    "User Interface Requirements": """
    Specify user interface requirements including usability standards, accessibility requirements, responsive design criteria, user experience guidelines, and interface design patterns. Include wireframes references, interaction flows, and visual design standards.
    """,

    "Integration Requirements": """
    Specify integration requirements including API specifications, data exchange formats, communication protocols, authentication mechanisms, and error handling procedures. Include integration patterns, service level agreements, and monitoring requirements.
    """,

    "Security Requirements": """
    Define comprehensive security requirements including authentication, authorization, data protection, network security, application security, and operational security. Include threat modeling, security controls, compliance requirements, and incident response procedures.
    """,
    
    "Product Future": """
    Give me Product Future in 1 paragraph of 4–5 lines highlighting roadmap and upcoming features.
    """,
    
    "Acceptance Criteria": """
    Define detailed acceptance criteria including functional acceptance criteria, performance acceptance criteria, security acceptance criteria, and operational acceptance criteria. Include test scenarios, success metrics, and sign-off procedures with measurable outcomes and validation methods.
    """,
    
        "Open Queries": """
    List Open Queries as a numbered list with each query in 1–2 lines.
    """
}

# Base IEEE 29148 formatting directives that apply to all headings
BASE_IEEE_DIRECTIVES = """
IEEE 29148 STYLE AND FORMATTING DIRECTIVES:
1. Write in precise, formal prose suitable for contractual and technical documentation.
2. Use present tense and the keyword "shall" for all mandatory requirements.
3. Begin content immediately under the section heading with no prefatory labels or commentary.
4. Use structured paragraphs and bullet points (•) for requirements or subpoints; avoid numbered lists such as 1.1, 1.2, etc.
5. Ensure every requirement is atomic, testable, and unambiguous; avoid vague or compound statements.
6. Maintain consistency in terminology (system, actor, module, interface) across all sections.
7. Do not reference the source of requirements (e.g., “meeting notes,” “uploaded document”); write as if original to the SRS.
8. Reference diagrams textually only as "Figure x.y" without embedding them.
9. Preserve a professional tone and avoid casual expressions.
10. Not put spacing between two bullet points in the SRS, formatted to appear as a polished deliverable.
11. Ensure structure follows the expected Table of Contents for SRS (Introduction, Scope, Stakeholders, Requirements, Architecture, NFRs, etc.).
12. Expand each section with sufficient depth to resemble an analyst’s authored SRS, not a summary or checklist.
"""

def get_heading_specific_prompt(heading: str) -> str:
    """
    Get the specific prompt for a given heading, or return a generic prompt if not found
    """
    return HEADING_SPECIFIC_PROMPTS.get(heading, """
    Generate comprehensive content for this section that aligns with professional SRS standards.
    Include detailed specifications, requirements, and implementation guidance appropriate for this section type.
    """)

def reset_diagram_counter():
    """Reset the global diagram counter for a new document"""
    global total_diagrams_generated, generated_diagram_signatures
    total_diagrams_generated = 0
    generated_diagram_signatures.clear()
    print("🔄 Reset diagram counter and signatures for new document")

# Define sections that should have diagrams
DIAGRAM_SECTIONS = {
    'sequence diagram': 'sequenceDiagram',
    'system architecture': 'flowchart',
    'architecture': 'flowchart',
    'system design': 'flowchart',
    'data flow': 'flowchart',
    'workflow': 'flowchart',
    'process flow': 'flowchart',
    'user interface': 'flowchart',
    'database design': 'erDiagram',
    'data model': 'erDiagram',
    'entity relationship': 'erDiagram',
    'class diagram': 'classDiagram',
    'component diagram': 'flowchart',
    'deployment': 'flowchart',
    'network': 'flowchart'
}

def should_generate_diagram(heading: str, user_prompt: str = "") -> str:
    """
    Check if a heading should have a diagram and return the diagram type
    Returns diagram type if should generate, None otherwise
    """
    # First check hardcoded section names
    heading_lower = heading.lower()
    for keyword, diagram_type in DIAGRAM_SECTIONS.items():
        if keyword in heading_lower:
            return diagram_type

    # If user prompt contains diagram-related keywords, detect diagram type from prompt
    if user_prompt and user_prompt.strip():
        prompt_lower = user_prompt.lower()

        # Check for diagram type keywords in user prompt (order matters - more specific first)
        if any(keyword in prompt_lower for keyword in ['sequence diagram', 'sequence', 'interaction']):
            return 'sequenceDiagram'
        elif any(keyword in prompt_lower for keyword in ['flowchart', 'architecture', 'system design', 'workflow', 'process', 'flow']):
            return 'flowchart'
        elif any(keyword in prompt_lower for keyword in ['er diagram', 'entity relationship', 'database schema', 'database design']):
            return 'erDiagram'
        elif any(keyword in prompt_lower for keyword in ['class diagram', 'uml', 'classes']):
            return 'classDiagram'
        elif 'diagram' in prompt_lower:
            # Default to flowchart if user mentions "diagram" but doesn't specify type
            return 'flowchart'

    return None

def fix_mermaid_syntax(mermaid_code: str) -> str:
    """
    Bulletproof post-processing cleanup for Gemini's Mermaid output
    Ensures mmdc CLI compatibility even when Gemini breaks rules
    """
    if not mermaid_code:
        return mermaid_code

    import re

    print(f"🔧 Applying bulletproof Mermaid cleanup...")

    # Step 1: Basic cleanup
    mermaid_code = mermaid_code.replace('\r\n', '\n').replace('\r', '\n').strip()

    # Step 1.1: CRITICAL - Remove any ellipsis from node labels
    # This is the most important fix to prevent "..." in diagram names
    mermaid_code = re.sub(r'\[([^[\]]*?)\.\.\.([^[\]]*?)\]', r'[\1 \2]', mermaid_code)
    mermaid_code = re.sub(r'\[([^[\]]*?)…([^[\]]*?)\]', r'[\1 \2]', mermaid_code)  # Unicode ellipsis
    print(f"🔧 Removed ellipsis from node labels")

    # Step 2: Critical fixes (based on your analysis)

    # Fix 1: Remove ALL single quotes and replace with double quotes
    mermaid_code = re.sub(r"subgraph\s+'([^']+)'", r'subgraph "\1"', mermaid_code)
    mermaid_code = mermaid_code.replace("'", '"')  # Replace any remaining single quotes

    # Fix 1.1: Remove/normalize HTML and special tags that mmdc cannot parse
    # Replace <br>, <br/>, <br /> with a space; strip any other HTML tags; normalize ampersands
    mermaid_code = re.sub(r'<br\s*/?>', ' ', mermaid_code, flags=re.IGNORECASE)
    mermaid_code = re.sub(r'<[^>]+>', ' ', mermaid_code)  # strip any remaining tags
    mermaid_code = mermaid_code.replace('&', 'and')

    # Fix 1.2: Handle problematic characters in node labels
    # Replace parentheses in node labels with spaces or remove them
    mermaid_code = re.sub(r'\[([^\]]*)\(([^)]*)\)([^\]]*)\]', r'[\1 \2 \3]', mermaid_code)
    # Remove forward slashes that can cause issues
    mermaid_code = re.sub(r'\[([^\]]*)/([^\]]*)\]', r'[\1 or \2]', mermaid_code)

    # Fix 2: Ensure newline after subgraph declarations (handle concatenation)
    # Pattern: subgraph "Name"A[Node] -> subgraph "Name"\n    A[Node]
    mermaid_code = re.sub(r'(subgraph\s+"[^"]+")([A-Za-z0-9])', r'\1\n    \2', mermaid_code)
    mermaid_code = re.sub(r'(subgraph\s+[A-Za-z0-9_]+)([A-Za-z0-9])', r'\1\n    \2', mermaid_code)

    # Fix 3: Separate concatenated nodes completely
    # Pattern: ][NodeID[ -> ]\n    NodeID[
    mermaid_code = re.sub(r'\]([A-Za-z0-9_]+)\[', r']\n    \1[', mermaid_code)

    # Step 3: Line-by-line processing for final cleanup
    lines = mermaid_code.split('\n')
    cleaned_lines = []
    diagram_type = None
    in_subgraph = False

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Detect diagram type
        if line.startswith(('flowchart', 'graph', 'sequenceDiagram', 'erDiagram', 'classDiagram')):
            if line.startswith('graph'):
                line = line.replace('graph', 'flowchart')
                diagram_type = 'flowchart'
            else:
                diagram_type = line.split()[0]
            cleaned_lines.append(line)
            continue

        # Track subgraph state
        if line.startswith('subgraph'):
            in_subgraph = True
            # Ensure proper quoting for subgraph names
            if not ('"' in line):
                # Extract name and add quotes
                parts = line.split(' ', 1)
                if len(parts) > 1:
                    name = parts[1].strip().replace('_', ' ')
                    line = f'subgraph "{name}"'
            cleaned_lines.append(line)
            continue
        elif line == 'end':
            in_subgraph = False
            cleaned_lines.append(line)
            continue

        # Fix arrows based on diagram type
        if diagram_type == 'flowchart':
            # Convert ALL arrow variants to --> (your recommendation)
            line = re.sub(r'([A-Za-z0-9_]+)\s*-+>\s*([A-Za-z0-9_]+)', r'\1 --> \2', line)

            # Ensure proper indentation in subgraphs
            if in_subgraph and not line.startswith('    '):
                line = '    ' + line

        elif diagram_type == 'sequenceDiagram':
            # Fix sequence diagram syntax issues

            # Fix 1: Handle malformed comma syntax like "Customer,Restaurant,CanarySystem: Call"
            if ',' in line and ':' in line and not line.startswith('participant'):
                # This is a malformed message line with commas
                parts = line.split(':', 1)
                if len(parts) == 2:
                    message = parts[1].strip()
                    participant_part = parts[0].strip()

                    if ',' in participant_part:
                        # Extract participants from comma-separated list
                        participants = [p.strip() for p in participant_part.split(',')]
                        if len(participants) >= 2:
                            # Create proper sequence message
                            line = f'    {participants[0]}->>+{participants[1]}: {message}'
                        else:
                            # Fallback to single participant
                            line = f'    {participants[0]}->>System: {message}'

            # Fix 2: Handle multiple participants declared with commas
            elif line.startswith('participant') and ',' in line:
                # Split multiple participants and create separate declarations
                participants_text = line.replace('participant', '').strip()
                participants = [p.strip() for p in participants_text.split(',')]
                # Skip this line, we'll add proper participant declarations later
                # For now, just clean it up
                if participants:
                    line = f'    participant {participants[0]}'

            # Fix 3: Remove any remaining commas from sequence diagrams
            elif ',' in line:
                line = line.replace(',', '')

            # Fix 4: Ensure proper sequence arrow syntax
            if '->' in line and not ('->>' in line or '-->' in line):
                line = line.replace('->', '->>')

        # Remove semicolons at end of lines
        if line.endswith(';'):
            line = line[:-1]

        cleaned_lines.append(line)

    result = '\n'.join(cleaned_lines)

    # Final cleanup
    result = re.sub(r'\n\s*\n\s*\n', '\n\n', result)  # Remove excessive newlines
    result = result.strip()

    print(f"🔧 Cleanup complete. Length: {len(mermaid_code)} -> {len(result)}")

    return result


def fix_mermaid_with_gemini(broken_code: str, error_message: str) -> str:
    """
    Ask Gemini to fix its own broken Mermaid code based on the error message
    """
    try:
        import google.generativeai as genai

        genai.configure(api_key="AIzaSyDS7AmlDi1cweQS1p-mJAUlB3uYHrYJXfI")
        model = genai.GenerativeModel('gemini-2.0-flash')

        prompt = f"""
        You generated this Mermaid code that has syntax errors:

        BROKEN CODE:
        {broken_code}

        ERROR MESSAGE:
        {error_message}

        CRITICAL ISSUE: The error shows broken text like "Cal" followed by "l Flow" on separate lines.
        This means the subgraph name got split incorrectly.

        SPECIFIC FIXES NEEDED:
        1. If you see broken text like "Cal" and "l Flow", combine them into "Call Flow"
        2. Subgraph names must be complete and properly quoted: subgraph "Call Flow"
        3. All content inside subgraphs must be indented with 4 spaces
        4. Each node must be on its own line
        5. Use --> for flowcharts (double arrows)
        6. Never use single quotes anywhere

        EXAMPLE OF CORRECT FORMAT:
        flowchart LR
            subgraph "Call Flow"
                A[Customer] --> B[Phone Number]
                B --> C[Conference]
            end

        Return ONLY the corrected Mermaid code, nothing else:
        """

        print(f"🔧 Asking Gemini to fix its own syntax errors...")
        response = model.generate_content(prompt)

        if response and response.text:
            fixed_code = response.text.strip()
            # Remove any markdown code blocks
            if fixed_code.startswith('```'):
                lines = fixed_code.split('\n')
                fixed_code = '\n'.join(lines[1:-1]) if len(lines) > 2 else fixed_code

            print(f"✅ Gemini provided syntax fix")
            return fixed_code
        else:
            print(f"❌ Gemini failed to provide fix")
            return broken_code

    except Exception as e:
        print(f"❌ Error asking Gemini for fix: {e}")
        return broken_code


def analyze_diagram_complexity(mermaid_code: str) -> dict:
    """
    Analyze if a Mermaid diagram is too complex for A4 page
    Returns analysis with suggestions for improvement
    """
    import re

    analysis = {
        'is_too_complex': False,
        'node_count': 0,
        'subgraph_count': 0,
        'connection_count': 0,
        'suggestions': []
    }

    if not mermaid_code:
        return analysis

    # Count nodes
    nodes = re.findall(r'([A-Za-z0-9_]+)\[([^\]]+)\]', mermaid_code)
    analysis['node_count'] = len(nodes)

    # Count subgraphs
    subgraphs = re.findall(r'subgraph\s+', mermaid_code)
    analysis['subgraph_count'] = len(subgraphs)

    # Count connections
    connections = re.findall(r'-->', mermaid_code)
    connections.extend(re.findall(r'->>', mermaid_code))
    analysis['connection_count'] = len(connections)

    # STRICT A4 COMPATIBILITY CHECKS
    # A4 portrait: 8.27" x 11.69" - diagrams must fit comfortably
    if analysis['node_count'] > 6:  # Reduced from 8 to 6
        analysis['is_too_complex'] = True
        analysis['suggestions'].append(f"Too many nodes ({analysis['node_count']}). Limit to 4-6 nodes per diagram for A4 compatibility.")

    if analysis['subgraph_count'] > 2:  # Reduced from 3 to 2
        analysis['is_too_complex'] = True
        analysis['suggestions'].append(f"Too many subgraphs ({analysis['subgraph_count']}). Limit to 1-2 subgraphs per diagram.")

    if analysis['connection_count'] > 8:  # Reduced from 12 to 8
        analysis['is_too_complex'] = True
        analysis['suggestions'].append(f"Too many connections ({analysis['connection_count']}). Simplify the flow to 6-8 connections maximum.")

    # Check for long node labels that make diagrams hard to read on A4
    long_labels = [label for _, label in nodes if len(label) > 15]  # Reduced from 20 to 15
    if long_labels:
        analysis['is_too_complex'] = True
        analysis['suggestions'].append("Some node labels are too long. Use 2-3 words maximum for A4 readability.")

            # Check for sequence diagram complexity
        if 'sequencediagram' in mermaid_code.lower():
            participant_count = len(re.findall(r'participant\s+', mermaid_code))
            if participant_count > 4:  # Limit sequence diagrams to 4 participants
                analysis['is_too_complex'] = True
                analysis['suggestions'].append(f"Too many participants ({participant_count}). Limit to 3-4 participants for A4 compatibility.")

            # Check sequence diagram message complexity
            message_count = len(re.findall(r'->>|-->>', mermaid_code))
            if message_count > 6:  # Limit messages for A4 compatibility
                analysis['is_too_complex'] = True
                analysis['suggestions'].append(f"Too many messages ({message_count}). Limit to 4-6 messages for A4 compatibility.")

    # Check for ER diagram complexity
    if 'erdiagram' in mermaid_code.lower():
        entity_count = len(re.findall(r'[A-Z_]+ \|\|', mermaid_code))
        if entity_count > 3:  # Limit ER diagrams to 3 entities
            analysis['is_too_complex'] = True
            analysis['suggestions'].append(f"Too many entities ({entity_count}). Limit to 2-3 entities for A4 compatibility.")

    return analysis


def generate_compact_diagram_prompt(heading: str, content: str, analysis: dict = None) -> str:
    """
    Generate a prompt specifically for compact A4-friendly diagrams
    """
    base_prompt = f"""
    Create a COMPACT Mermaid diagram for "{heading}" that fits perfectly on an A4 portrait page.

    CONTENT CONTEXT:
    {content[:1000]}...

    STRICT SIZE REQUIREMENTS:
    - Maximum 6-8 nodes total
    - Maximum 2-3 subgraphs
    - Use concise labels (3-4 words max)
    - Prefer vertical layout (TD direction)
    - Focus on CORE components only

    If the system is complex, create a HIGH-LEVEL overview diagram showing only the most important components.
    Do NOT try to include every detail - prioritize clarity and readability on A4 page.
    """

# --- Coverage helpers -------------------------------------------------------

def extract_components_from_text(text: str) -> list:
    """Extract likely system components from SRS text for coverage analysis."""
    import re
    tokens = set()
    if not text:
        return []
    t = text
    # Common domain/system words and technologies
    patterns = [
        r"\b(API Gateway|Auth Service|Authentication Service|Authorization Service|Notification Service)\b",
        r"\b(Microservice|Service|Module|Component|Subsystem|Worker|Scheduler)\b",
        r"\b(Database|DB|PostgreSQL|MySQL|MongoDB|Redis|Kafka|Queue|Broker)\b",
        r"\b(Frontend|Web App|Mobile App|Client|Backend|Server)\b",
        r"\b(Logging|Monitoring|Tracing|Metrics|Observability)\b",
        r"\b(ETL|Ingestion|Processing|Analytics|Reporting)\b",
        r"\b(Identity Provider|IdP|SSO|OAuth|OIDC)\b"
    ]
    for p in patterns:
        for m in re.findall(p, t, flags=re.IGNORECASE):
            tokens.add(m.strip())
    # Capitalized multi-word terms
    for m in re.findall(r"\b([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){0,3})\b", t):
        if len(m) > 2 and not m.isupper():
            tokens.add(m.strip())
    return list(tokens)[:30]


def extract_names_from_mermaid(code: str) -> list:
    """Extract participant/node/subgraph names from Mermaid code."""
    import re
    parts = set()
    if not code:
        return []
    parts.update(re.findall(r'^\s*participant\s+"?([A-Za-z0-9 _-]+)"?', code, flags=re.MULTILINE))
    parts.update(re.findall(r'\[([^\]]+)\]', code))
    parts.update(re.findall(r'subgraph\s+"([^"]+)"', code))
    cleaned = []
    for p in parts:
        x = re.sub(r'[^A-Za-z0-9 ]+', ' ', p).strip()
        if x and len(x) > 1:
            cleaned.append(x)
    return cleaned[:30]


    if analysis and analysis.get('suggestions'):
        base_prompt += f"\n\nPREVIOUS ISSUES TO AVOID:\n" + "\n".join(f"- {s}" for s in analysis['suggestions'])

    return base_prompt


def generate_sub_diagrams(heading: str, content: str, diagram_type: str, model, original_diagram: str) -> list:
    """
    Generate multiple focused sub-diagrams using AI analysis - completely AI-driven approach
    """
    try:
        print(f"🔄 Using AI to break complex diagram into unique sub-diagrams...")

        # AI-driven validation and planning for necessary diagrams only
        # Also ensure coverage of the WHOLE system by focusing on components not covered in the original diagram
        def _extract_names_from_mermaid(code: str) -> list:
            import re
            parts = set()
            parts.update(re.findall(r'^\s*participant\s+"?([A-Za-z0-9 _-]+)"?', code, flags=re.MULTILINE))
            parts.update(re.findall(r'\[([^\]]+)\]', code))
            parts.update(re.findall(r'subgraph\s+"([^"]+)"', code))
            cleaned = []
            for p in parts:
                t = re.sub(r'[^A-Za-z0-9 ]+', ' ', p).strip()
                if t and len(t) > 1:
                    cleaned.append(t)
            return cleaned[:20]

        covered_components = _extract_names_from_mermaid(original_diagram)

        validation_prompt = f"""
        Analyze this system content and determine if it actually needs to be split into sub-diagrams.

        SYSTEM HEADING: {heading}
        DIAGRAM TYPE: {diagram_type}
        SYSTEM CONTENT: {content[:1200]}
        ORIGINAL DIAGRAM (ALREADY COVERED COMPONENTS): {', '.join(covered_components)}

        GOAL:
        - Ensure the entire SRS/system is covered across the main diagram + sub-diagrams
        - Sub-diagrams must focus on components/flows NOT covered in the original diagram

        VALIDATION REQUIREMENTS:
        1. Only suggest sub-diagrams if the system is genuinely complex and has distinct aspects
        2. Don't add unnecessary diagrams - validate if splitting is actually beneficial
        3. Maximum 2 sub-diagrams (not 3) - only if truly needed
        4. Each diagram must be based on actual SRS content, not generic assumptions
        5. Avoid duplicate or similar diagrams

        Your task:
        - First determine: Does this system actually need sub-diagrams? (YES/NO)
        - If YES, create 1-2 unique sub-diagram plans based on actual SRS content
        - Ensure the plans primarily cover UN-COVERED components/flows (not in: {', '.join(covered_components)})
        - If NO, return empty array

        Return ONLY a JSON response like this:
        {{
            "needs_sub_diagrams": true/false,
            "reason": "Brief explanation why sub-diagrams are/aren't needed",
            "sub_diagrams": [
                {{
                    "focus_area": "Actual system aspect from SRS content",
                    "required_components": ["ActualComponent1", "ActualComponent2"],
                    "forbidden_components": ["ActualComponent3", "ActualComponent4"],
                    "specific_instructions": "Instructions based on actual SRS content"
                }}
            ]
        }}

        CRITICAL: Only generate diagrams that are actually needed and based on real SRS content.
        """

        try:
            validation_response = model.generate_content(validation_prompt)
            if validation_response and validation_response.text:
                import json
                import re

                # Extract JSON from response
                json_text = validation_response.text.strip()
                json_text = re.sub(r'```json\s*', '', json_text)
                json_text = re.sub(r'```\s*$', '', json_text)

                try:
                    validation_result = json.loads(json_text)
                    needs_sub_diagrams = validation_result.get('needs_sub_diagrams', False)
                    reason = validation_result.get('reason', 'No reason provided')

                    print(f"🔍 AI validation: Needs sub-diagrams = {needs_sub_diagrams}")
                    print(f"📝 Reason: {reason}")

                    if needs_sub_diagrams:
                        sub_diagram_plan = validation_result.get('sub_diagrams', [])
                        print(f"✅ AI generated {len(sub_diagram_plan)} necessary sub-diagram plans")
                    else:
                        print("✅ AI determined sub-diagrams are not necessary - using original diagram")
                        return []  # No sub-diagrams needed

                except json.JSONDecodeError:
                    print("⚠️ AI validation response was not valid JSON, checking if sub-diagrams are needed")
                    # Simple fallback validation
                    if len(content) < 500:
                        print("✅ Content is simple, no sub-diagrams needed")
                        return []
                    else:
                        print("⚠️ Using minimal fallback sub-diagrams")
                        sub_diagram_plan = [
                            {
                                "focus_area": "Core System Flow",
                                "required_components": ["Main", "Process"],
                                "forbidden_components": ["External"],
                                "specific_instructions": "Show core system interactions only"
                            }
                        ]
            else:
                print("⚠️ No AI validation response, checking content complexity")
                if len(content) < 500:
                    print("✅ Content is simple, no sub-diagrams needed")
                    return []
                else:
                    print("⚠️ Using minimal fallback")
                    sub_diagram_plan = [
                        {
                            "focus_area": "Core System Flow",
                            "required_components": ["Main", "Process"],
                            "forbidden_components": ["External"],
                            "specific_instructions": "Show core system interactions only"
                        }
                    ]
        except Exception as e:
            print(f"⚠️ AI validation failed: {e}, checking content complexity")
            if len(content) < 500:
                print("✅ Content is simple, no sub-diagrams needed")
                return []
            else:
                print("⚠️ Using minimal fallback")
                sub_diagram_plan = [
                    {
                        "focus_area": "Core System Flow",
                        "required_components": ["Main", "Process"],
                        "forbidden_components": ["External"],
                        "specific_instructions": "Show core system interactions only"
                    }
                ]

        def enhanced_signature_for(code: str) -> str:
            """SOLUTION: Enhanced signature that detects flow patterns, not just labels"""
            import re
            code = code.lower()

            # Extract all components
            labels = re.findall(r'\[([^\]]+)\]', code)
            subgraphs = re.findall(r'subgraph\s+"([^"]+)"', code)
            participants = re.findall(r'^\s*participant\s+"?([A-Za-z0-9 _-]+)"?', code, flags=re.MULTILINE)

            # Extract flow patterns (more sophisticated)
            edges = re.findall(r'([a-z0-9_]+)\s*-+>+\s*([a-z0-9_]+)', code)
            sequence_flows = re.findall(r'([a-z0-9_]+)->>([a-z0-9_]+)', code)

            # Create comprehensive signature
            tokens = set()

            # Add normalized component names
            for s in labels + subgraphs + participants:
                t = re.sub(r'[^a-z0-9 ]+', ' ', s)
                t = re.sub(r'\s+', ' ', t).strip()
                if t and len(t) > 1:
                    tokens.add(f"comp:{t}")

            # Add flow patterns with direction
            for a, b in edges + sequence_flows:
                tokens.add(f"flow:{a}->{b}")

            # Add structural patterns
            if 'sequencediagram' in code:
                tokens.add("type:sequence")
                # Count interaction patterns
                interactions = len(re.findall(r'->>|-->', code))
                tokens.add(f"interactions:{interactions}")
            elif 'flowchart' in code:
                tokens.add("type:flowchart")
                # Count node patterns
                nodes = len(re.findall(r'\[[^\]]+\]', code))
                tokens.add(f"nodes:{nodes}")

            # Add complexity indicators
            if len(edges) > 5:
                tokens.add("complex:high")
            elif len(edges) > 2:
                tokens.add("complex:medium")
            else:
                tokens.add("complex:low")

            return '|'.join(sorted(tokens))

        seen_signatures = set()
        sub_diagrams = []

        for i, sub_plan in enumerate(sub_diagram_plan, 1):
            focus_area = sub_plan.get('focus_area', f'Sub-diagram {i}')
            required_components = sub_plan.get('required_components', [])
            forbidden_components = sub_plan.get('forbidden_components', [])
            specific_instructions = sub_plan.get('specific_instructions', '')

            # AI-driven prompt with forced differentiation and STRICT A4 requirements
            sub_prompt = f"""
            Generate a Mermaid {diagram_type} for this SPECIFIC system aspect:

            SYSTEM: {heading}
            FOCUS AREA: {focus_area}

            SYSTEM CONTEXT:
            {content[:800]}

            AI-GENERATED COMPONENT REQUIREMENTS:
            MUST USE ONLY: {', '.join(required_components)}
            ABSOLUTELY FORBIDDEN: {', '.join(forbidden_components)}

            AI-GENERATED INSTRUCTIONS:
            {specific_instructions}

            STRICT A4 COMPATIBILITY REQUIREMENTS:
            1. Diagram MUST fit on A4 portrait page (8.27" x 11.69")
            2. Maximum 4-6 nodes/participants ONLY
            3. Maximum 1-2 subgraphs ONLY
            4. Maximum 6-8 connections ONLY
            5. Node labels must be 2-3 words maximum
            6. Use vertical layout (TD direction) for better A4 fit
            7. Focus on CORE components only, not every detail

            CRITICAL NAMING RULES (MUST FOLLOW EXACTLY):
            1. Every node must have a meaningful, descriptive, and complete label.
            2. Never output single-letter node names (e.g., "e" or "a").
            Replace them with full system terms like "Engine", "Endpoint", or "Entity".
            3. Node labels must use full words, ABSOLUTELY NO truncation or ellipsis (...) ANYWHERE.
            4. Keep node names short (max 3 words) but fully descriptive.
            Example: "Data Storage", "Reporting Service", "User Interface".
            5. Avoid special characters, parentheses, commas, and periods in node labels.
            6. Always use Title Case (capitalize first letter of each word).
            7. If system entities are unclear, use safe defaults: "Frontend", "Backend", "Database", "Reporting Service", "API Gateway".
            8. If a label is too long, shorten it by choosing the 2–3 most important words, NEVER use ellipsis (...).
            9. FORBIDDEN: Any use of "..." or ellipsis in node names - this will cause diagram generation to fail.
            10. REQUIRED: All node labels must be complete, readable words without any truncation symbols.


            STRICT ENFORCEMENT RULES:
            1. You MUST use ONLY the components listed in "MUST USE ONLY"
            2. You are ABSOLUTELY FORBIDDEN from using any components in "ABSOLUTELY FORBIDDEN"
            3. If you use any forbidden components, the diagram will be rejected
            4. Each diagram must show completely different participants/nodes
            5. Each diagram must show completely different flow patterns
            6. PRIORITIZE A4 COMPATIBILITY over completeness

            DIAGRAM REQUIREMENTS:
            - Generate Mermaid diagrams for the given SRS sections
            - Each diagram must be compact enough to fit within a single A4 portrait page and must be readable (not blurry)
            - Each sub-diagram must be unique (no repetition of the same diagram text or flow)
            - Do not add headings, titles, or duplicate captions outside the diagrams
            - Use short labels (max 2–3 words) for each node
            - Keep the style minimal, clean, and hierarchical so that the diagrams are clear when rendered in DOCX/PDF
            - Maximum 4-6 nodes/participants only
            - Use {diagram_type} format
            - This is part {i} of {len(sub_diagram_plan)} - must show COMPLETELY DIFFERENT components/flows than other parts

            VALIDATION CHECK:
            Before generating, verify:
            - Are you using ONLY the required components?
            - Are you avoiding ALL forbidden components?
            - Is this flow pattern different from typical generic flows?

            CRITICAL MERMAID SYNTAX RULES:
            1. Never use single quotes anywhere
            2. For subgraphs with spaces: subgraph "Name"
            3. Use --> for flowcharts, ->> for sequence diagrams
            4. Keep node labels short and clean
            5. No semicolons, commas, or special characters in labels

            Return ONLY the Mermaid code for this single diagram (no code blocks, no multiple diagrams):
            """

            try:
                response = model.generate_content(sub_prompt)
                if response and response.text:
                    sub_code = response.text.strip()
                    if sub_code.startswith('```'):
                        lines = sub_code.split('\n')
                        sub_code = '\n'.join(lines[1:-1]) if len(lines) > 2 else sub_code

                    # Apply syntax fixes
                    sub_code = fix_mermaid_syntax(sub_code)

                    # SOLUTION: Validate single diagram per response
                    diagram_count = sub_code.lower().count('sequencediagram') + sub_code.lower().count('flowchart') + sub_code.lower().count('erdiagram') + sub_code.lower().count('classdiagram')
                    if diagram_count > 1:
                        print(f"⚠️ Sub-diagram {i} contains multiple diagrams ({diagram_count}), extracting first one only")
                        # Extract only the first diagram
                        lines = sub_code.split('\n')
                        first_diagram_lines = []
                        diagram_started = False
                        for line in lines:
                            if any(keyword in line.lower() for keyword in ['sequencediagram', 'flowchart', 'erdiagram', 'classdiagram']):
                                if diagram_started:
                                    break  # Stop at second diagram
                                diagram_started = True
                            if diagram_started:
                                first_diagram_lines.append(line)
                        sub_code = '\n'.join(first_diagram_lines)
                        print(f"✅ Extracted single diagram from response")

                    # Check for duplicate diagrams using enhanced signature
                    diagram_signature = enhanced_signature_for(sub_code)
                    if diagram_signature in seen_signatures:
                        print(f"⚠️ Sub-diagram {i} is duplicate (signature: {diagram_signature[:50]}...), skipping")
                        continue

                    # Additional check: if signature is too similar to existing ones (>80% overlap)
                    for existing_sig in seen_signatures:
                        existing_tokens = set(existing_sig.split('|'))
                        new_tokens = set(diagram_signature.split('|'))
                        if existing_tokens and new_tokens:
                            overlap = len(existing_tokens.intersection(new_tokens)) / len(existing_tokens.union(new_tokens))
                            if overlap > 0.8:
                                print(f"⚠️ Sub-diagram {i} is too similar to existing diagram ({overlap:.1%} overlap), skipping")
                                continue

                    # Verify this sub-diagram is actually simpler and A4 compatible
                    sub_complexity = analyze_diagram_complexity(sub_code)
                    if not sub_complexity['is_too_complex']:
                        # Add signature to seen set to prevent future duplicates
                        seen_signatures.add(diagram_signature)
                        sub_diagrams.append({
                            'title': f"{heading} - Part {i}: {focus_area}",
                            'mermaid_code': sub_code,
                            'focus_area': focus_area,
                            'part_number': i
                        })
                        print(f"✅ Generated sub-diagram {i}: {focus_area} ({sub_complexity['node_count']} nodes)")
                        print(f"📝 Diagram signature: {diagram_signature[:50]}...")
                    else:
                        print(f"⚠️ Sub-diagram {i} still too complex for A4, skipping")
                        print(f"📋 Issues: {', '.join(sub_complexity['suggestions'])}")

            except Exception as sub_error:
                print(f"❌ Failed to generate sub-diagram {i}: {sub_error}")

        # If we have fewer than 2 unique sub-diagrams, the original was probably not that complex
        if len(sub_diagrams) < 2:
            print(f"⚠️ Only generated {len(sub_diagrams)} unique sub-diagrams, original diagram may not be too complex")
            return []

        return sub_diagrams

    except Exception as e:
        print(f"❌ Error generating sub-diagrams: {e}")
        return []


def insert_sub_diagrams_into_document(doc: Document, heading: str, sub_diagrams: list):
    """
    Insert multiple sub-diagrams into the document for a complex section
    """
    try:
        print(f"📄 Inserting {len(sub_diagrams)} sub-diagrams for: {heading}")

        # Find the section in the document
        target_paragraph = None
        for para in doc.paragraphs:
            if heading.lower() in para.text.lower() and para.style.name == 'SRS Heading 2':
                target_paragraph = para
                break

        if not target_paragraph:
            print(f"⚠️ Could not find section '{heading}' in document")
            return

        # Insert each sub-diagram without subheadings
        for i, sub_diagram in enumerate(sub_diagrams):
            # Generate PNG for this sub-diagram (no titles or descriptions)
            sub_diagram_filename = f"{heading.replace(' ', '_')}_Part_{sub_diagram['part_number']}_diagram.png"
            sub_diagram_path = os.path.join(os.path.dirname(__file__), '..', 'temp_diagrams', sub_diagram_filename)

            # Ensure sub-diagram code is properly cleaned before conversion
            cleaned_sub_code = fix_mermaid_syntax(sub_diagram['mermaid_code'])

            # Convert to PNG
            if convert_mermaid_to_png_with_fallback(cleaned_sub_code, sub_diagram_path, sub_diagram['title']):
                # Insert the PNG
                diagram_para = doc.add_paragraph()
                run = diagram_para.runs[0] if diagram_para.runs else diagram_para.add_run()
                run.add_picture(sub_diagram_path, width=Inches(6))
                diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"✅ Inserted sub-diagram {i+1}: {sub_diagram['title']}")
            else:
                # Skip failed diagrams (no placeholder text since we don't want subheadings)
                print(f"❌ Failed to generate PNG for sub-diagram {i+1}, skipping")

        print(f"📄 Completed inserting {len(sub_diagrams)} sub-diagrams")

    except Exception as e:
        print(f"❌ Error inserting sub-diagrams: {e}")


def convert_mermaid_to_png_with_error_capture(mermaid_code: str, output_path: str) -> tuple[bool, str]:
    """
    Convert Mermaid to PNG and capture error messages for Gemini correction
    Returns: (success: bool, error_message: str)
    """
    try:
        mmdc_cmd = find_mmdc_command()
        if not mmdc_cmd:
            return False, "mmdc command not found"

        # Create temporary file for mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as temp_file:
            temp_file.write(mermaid_code)
            temp_mmd_path = temp_file.name

        print(f"📝 Temp Mermaid file: {temp_mmd_path}")
        print(f"📋 Mermaid code preview: {mermaid_code[:100]}...")

        # Build command for PNG
        if mmdc_cmd.startswith('npx'):
            cmd = mmdc_cmd.split() + [
                '-i', temp_mmd_path,
                '-o', output_path,
                '-b', 'white',
            ]
        else:
            cmd = [
                mmdc_cmd,
                '-i', temp_mmd_path,
                '-o', output_path,
                '-b', 'white',
            ]

        print(f"🚀 Running command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        # Clean up temporary file
        if temp_mmd_path and os.path.exists(temp_mmd_path):
            os.unlink(temp_mmd_path)

        print(f"📊 Command exit code: {result.returncode}")
        if result.stdout:
            print(f"📤 stdout: {result.stdout}")
        if result.stderr:
            print(f"📥 stderr: {result.stderr}")

        if result.returncode == 0:
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"✅ Successfully generated PNG: {output_path} ({file_size} bytes)")
                return True, ""
            else:
                return False, "PNG file was not created"
        else:
            error_msg = result.stderr or result.stdout or "Unknown error"
            print(f"❌ mmdc command failed with exit code {result.returncode}")
            print(f"   Error details: {error_msg}")
            return False, error_msg

    except subprocess.TimeoutExpired:
        error_msg = "mmdc command timed out"
        print(f"⏰ {error_msg}")
        return False, error_msg
    except Exception as e:
        error_msg = f"Exception during PNG conversion: {str(e)}"
        print(f"❌ {error_msg}")
        return False, error_msg


def generate_mermaid_diagram(heading: str, content: str, uploaded_content: str = "", user_prompt: str = "", full_srs_content: str = "") -> str:
    """
    Generate Mermaid diagram code using Gemini AI based on the heading and content
    Maximum 3 unique diagrams total per document with strict A4 compatibility
    """
    try:
        global total_diagrams_generated, generated_diagram_signatures

        # Check if we've already generated maximum diagrams
        if total_diagrams_generated >= 3:
            print(f"⚠️ Maximum 3 diagrams already generated, skipping diagram for: {heading}")
            return None

        print(f"🎨 Attempting to generate diagram {total_diagrams_generated + 1}/3 for: {heading}")

        # Configure Gemini API with error handling
        try:
            genai.configure(api_key="AIzaSyDS7AmlDi1cweQS1p-mJAUlB3uYHrYJXfI")
            model = genai.GenerativeModel('gemini-2.0-flash')
            print("✅ Gemini API configured successfully")
        except Exception as config_error:
            print(f"❌ Failed to configure Gemini API: {config_error}")
            return None

        diagram_type = should_generate_diagram(heading, user_prompt)
        print(f"📊 Diagram type for '{heading}': {diagram_type}")

        if not diagram_type:
            print(f"⚠️ No diagram type found for heading: {heading}")
            return None

        # ENHANCED PROMPT WITH STRICT A4 REQUIREMENTS
        if user_prompt and user_prompt.strip():
            # Process user prompt to ensure professional SRS language
            professional_prompt = user_prompt.replace("according to meeting summary", "based on system requirements")
            professional_prompt = professional_prompt.replace("meeting transcript", "system specifications")
            professional_prompt = professional_prompt.replace("meeting", "requirements analysis")

            prompt = f"""
            PROFESSIONAL SRS DIAGRAM REQUIREMENTS:
            {professional_prompt}

            SECTION HEADING: {heading}
            SECTION CONTENT: {content}
            SOURCE OF TRUTH: Full SRS (ignore any meeting/transcript inputs)

            FULL SRS DOCUMENT CONTEXT:
            {full_srs_content if full_srs_content else content}

            DIAGRAM TYPE TO GENERATE: {diagram_type}

            CRITICAL A4 COMPATIBILITY REQUIREMENTS:
            1. Diagram MUST fit on A4 portrait page (8.27" x 11.69")
            2. Maximum 4-6 nodes/participants ONLY
            3. Maximum 1-2 subgraphs ONLY
            4. Maximum 6-8 connections ONLY
            5. Node labels must be 2-3 words maximum
            6. Use vertical layout (TD direction) for better A4 fit
            7. Focus on CORE components only, not every detail
            8. All the nodes should contain proper & full namings

            CRITICAL NAMING RULES (MUST FOLLOW EXACTLY):
            1. Every node must have a meaningful, descriptive, and complete label.
            2. Never output single-letter node names (e.g., "e" or "a").
            Replace them with full system terms like "Engine", "Endpoint", or "Entity".
            3. Node labels must use full words, ABSOLUTELY NO truncation or ellipsis (...) ANYWHERE.
            4. Keep node names short (max 3 words) but fully descriptive.
            Example: "Data Storage", "Reporting Service", "User Interface".
            5. Avoid special characters, parentheses, commas, and periods in node labels.
            6. Always use Title Case (capitalize first letter of each word).
            7. If system entities are unclear, use safe defaults: "Frontend", "Backend", "Database", "Reporting Service", "API Gateway".
            8. If a label is too long, shorten it by choosing the 2–3 most important words, NEVER use ellipsis (...).
            9. FORBIDDEN: Any use of "..." or ellipsis in node names - this will cause diagram generation to fail.
            10. REQUIRED: All node labels must be complete, readable words without any truncation symbols.

            INSTRUCTIONS:
            - Base the diagram on the COMPLETE system described in the full SRS document
            - Show overall system architecture, workflows, and component interactions
            - Include all major components mentioned across the entire SRS
            - Focus on the big picture rather than just this section
            - PRIORITIZE A4 COMPATIBILITY over completeness

            TECHNICAL REQUIREMENTS:
            1. Return ONLY one single Mermaid diagram (no multiple diagrams, no code blocks)
            2. Use proper Mermaid {diagram_type} syntax
            3. CRITICAL: Avoid parentheses, commas, periods, and special characters in node labels
            4. Use simple node names like [Frontend], [Backend], [Database], [User], [API]
            5. Keep node text short and clean
            6. For ER diagrams: CRITICAL - Keep it VERY simple with max 3 entities only
            7. For ER diagrams: Use only basic data types (int, string, date) - NO complex types
            8. For ER diagrams: Use simple entity names like USER, ORDER, PRODUCT (no spaces or special chars)
            9. For ER diagrams: Maximum 3-4 fields per entity to ensure compatibility

            CRITICAL MERMAID SYNTAX RULES (MUST FOLLOW EXACTLY):
            1. Never use single quotes in any part of the Mermaid code.
            2. For subgraphs: If the name has spaces, wrap it in double quotes "Name".
            3. After subgraph declaration, always start a new line before listing nodes.
            4. Indent all nodes inside a subgraph with exactly 4 spaces.
            5. Each node definition must be on its own separate line.
            6. For flowcharts, use --> (double hyphen and angle bracket) for all connections.
            7. Never use ->, --, or other arrow variants.
            8. Avoid parentheses, commas, periods, or special characters in node IDs and labels.
            9. No semicolons after diagram type declarations.
            10. For sequence diagrams: NO COMMAS anywhere in the code.
            11. For sequence diagrams: Each participant must be declared separately.
            12. For sequence diagrams: Use ->> for messages, -->> for responses.
            13. Keep diagram code clean and minimal so mmdc can parse it.

            A4 COMPATIBILITY ENFORCEMENT:
            14. Generate compact, unique diagrams that fit within one A4 portrait page
            15. Avoid repetition of diagrams and flows - each diagram must be completely unique
            16. If a diagram is too large, break it into multiple smaller sub-diagrams with different flows
            17. Ensure sub-diagrams are unique and represent different system aspects
            18. Do not add headings, titles, or duplicate captions outside the diagrams
            19. Use short labels (max 2–3 words) for each node
            20. Keep the style minimal, clean, and hierarchical for clear rendering in DOCX/PDF
            21. Maximum 4-6 nodes/participants per diagram for A4 compatibility
            22. Maintain clear diagram hierarchy with no duplicate node names across diagrams
            23. Return ONLY one single Mermaid diagram (no multiple diagrams in one response)

            CORRECT MERMAID SYNTAX EXAMPLES FOR A4:

            For flowchart (COMPACT):
            flowchart TD
                A[Start] --> B[Process]
                B --> C[End]

            For flowchart with subgraphs (A4 COMPATIBLE):
            flowchart TD
                subgraph "Core System"
                    A[User Input]
                    B[Process]
                end
                A --> B --> C[Output]

            For sequenceDiagram (A4 COMPATIBLE):
            sequenceDiagram
                participant User
                participant System
                User->>System: Request
                System-->>User: Response

            For erDiagram (A4 COMPATIBLE):
            erDiagram
                USER ||--o{{ ORDER : places

            Generate the Mermaid diagram code based on the COMPLETE SRS system:
            """
        else:
            prompt = f"""
            Generate a Mermaid diagram for the complete system described in the SRS document.

            SECTION HEADING: {heading}
            SECTION CONTENT: {content}
            SOURCE OF TRUTH: Full SRS (ignore meeting summaries)

            FULL SRS DOCUMENT CONTEXT:
            {full_srs_content if full_srs_content else content}

            DIAGRAM TYPE: {diagram_type}

            CRITICAL A4 COMPATIBILITY REQUIREMENTS:
            1. Diagram MUST fit on A4 portrait page (8.27" x 11.69")
            2. Maximum 4-6 nodes/participants ONLY
            3. Maximum 1-2 subgraphs ONLY
            4. Maximum 6-8 connections ONLY
            5. Node labels must be 2-3 words maximum
            6. Use vertical layout (TD direction) for better A4 fit
            7. Focus on CORE components only, not every detail
            8. All the nodes should contain proper & full namings

            REQUIREMENTS:
            1. Create a professional {diagram_type} diagram showing the COMPLETE system architecture
            2. Base the diagram on the ENTIRE SRS document, not just this section
            3. Show all major components, their relationships, and interactions
            4. Include system workflows and data flows from the complete SRS
            5. Use clear, descriptive labels for all system components
            6. Return ONLY one single Mermaid diagram (no multiple diagrams, no code blocks)
            7. CRITICAL: Avoid parentheses, commas, periods, and special characters in node labels
            8. Use simple node names like [Frontend], [Backend], [Database], [User], [API], [Cache]
            9. Keep node text short and clean
            10. For ER diagrams: CRITICAL - Keep it VERY simple with max 3 entities only
            11. For ER diagrams: Use only basic data types (int, string, date) - NO complex types
            12. For ER diagrams: Use simple entity names like USER, ORDER, PRODUCT (no spaces or special chars)
            13. For ER diagrams: Maximum 3-4 fields per entity to ensure compatibility

            CRITICAL MERMAID SYNTAX RULES (MUST FOLLOW EXACTLY):
            1. Never use single quotes in any part of the Mermaid code.
            2. For subgraphs: If the name has spaces, wrap it in double quotes "Name".
            3. After subgraph declaration, always start a new line before listing nodes.
            4. Indent all nodes inside a subgraph with exactly 4 spaces.
            5. Each node definition must be on its own separate line.
            6. For flowcharts, use --> (double hyphen and angle bracket) for all connections.
            7. Never use ->, --, or other arrow variants.
            8. Avoid parentheses, commas, periods, or special characters in node IDs and labels.
            9. No semicolons after diagram type declarations.
            10. For sequence diagrams: NO COMMAS anywhere in the code.
            11. For sequence diagrams: Each participant must be declared separately.
            12. For sequence diagrams: Use ->> for messages, -->> for responses.
            13. Keep diagram code clean and minimal so mmdc can parse it.

            CRITICAL NAMING RULES (MUST FOLLOW EXACTLY):
            1. Every node must have a meaningful, descriptive, and complete label.
            2. Never output single-letter node names (e.g., "e" or "a").
            Replace them with full system terms like "Engine", "Endpoint", or "Entity".
            3. Node labels must use full words, ABSOLUTELY NO truncation or ellipsis (...) ANYWHERE.
            4. Keep node names short (max 3 words) but fully descriptive.
            Example: "Data Storage", "Reporting Service", "User Interface".
            5. Avoid special characters, parentheses, commas, and periods in node labels.
            6. Always use Title Case (capitalize first letter of each word).
            7. If system entities are unclear, use safe defaults: "Frontend", "Backend", "Database", "Reporting Service", "API Gateway".
            8. If a label is too long, shorten it by choosing the 2–3 most important words, NEVER use ellipsis (...).
            9. FORBIDDEN: Any use of "..." or ellipsis in node names - this will cause diagram generation to fail.
            10. REQUIRED: All node labels must be complete, readable words without any truncation symbols.

            A4 COMPATIBILITY ENFORCEMENT:
            14. Generate compact, unique diagrams that fit within one A4 portrait page
            15. Avoid repetition of diagrams and flows - each diagram must be completely unique
            16. If a diagram is too large, break it into multiple smaller sub-diagrams with different flows
            17. Ensure sub-diagrams are unique and represent different system aspects
            18. Do not add headings, titles, or duplicate captions outside the diagrams
            19. Use short labels (max 2–3 words) for each node
            20. Keep the style minimal, clean, and hierarchical for clear rendering in DOCX/PDF
            21. Maximum 4-6 nodes/participants per diagram for A4 compatibility
            22. Maintain clear diagram hierarchy with no duplicate node names across diagrams
            23. Return ONLY one single Mermaid diagram (no multiple diagrams in one response)

            CORRECT MERMAID SYNTAX EXAMPLES FOR A4:

            For flowchart (COMPACT):
            flowchart TD
                A[Start] --> B[Process]
                B --> C[End]

            For flowchart with subgraphs (A4 COMPATIBLE):
            flowchart TD
                subgraph "Core System"
                    A[User Input]
                    B[Process]
                end
                A --> B --> C[Output]

            For sequenceDiagram (A4 COMPATIBLE):
            sequenceDiagram
                participant User
                participant System
                User->>System: Request
                System-->>User: Response

            For erDiagram (A4 COMPATIBLE):
            erDiagram
                USER ||--o{{ ORDER : places

            Generate the Mermaid diagram code representing the complete system:
            """

        print(f"📤 Sending diagram prompt to Gemini for '{heading}'...")
        response = model.generate_content(prompt)

        if response and response.text:
            mermaid_code = response.text.strip()

            # Clean the mermaid code - remove markdown code fences if present
            if mermaid_code.startswith('```mermaid'):
                # Remove opening ```mermaid
                mermaid_code = mermaid_code[10:].strip()
            elif mermaid_code.startswith('```'):
                # Remove opening ```
                mermaid_code = mermaid_code[3:].strip()

            if mermaid_code.endswith('```'):
                # Remove closing ```
                mermaid_code = mermaid_code[:-3].strip()

            # SOLUTION: Validate single diagram per response
            diagram_count = mermaid_code.lower().count('sequencediagram') + mermaid_code.lower().count('flowchart') + mermaid_code.lower().count('erdiagram') + mermaid_code.lower().count('classdiagram')
            if diagram_count > 1:
                print(f"⚠️ Main diagram contains multiple diagrams ({diagram_count}), extracting first one only")
                # Extract only the first diagram
                lines = mermaid_code.split('\n')
                first_diagram_lines = []
                diagram_started = False
                for line in lines:
                    if any(keyword in line.lower() for keyword in ['sequencediagram', 'flowchart', 'erdiagram', 'classdiagram']):
                        if diagram_started:
                            break  # Stop at second diagram
                        diagram_started = True
                    if diagram_started:
                        first_diagram_lines.append(line)
                mermaid_code = '\n'.join(first_diagram_lines)
                print(f"✅ Extracted single diagram from main response")

            # Apply syntax fixes
            print(f"🔧 Applying syntax fixes to diagram...")
            fixed_mermaid_code = fix_mermaid_syntax(mermaid_code)

            # VALIDATE DIAGRAM FOR A4 COMPATIBILITY
            is_valid, error_msg, complexity = validate_diagram_for_a4(fixed_mermaid_code, heading)
            if not is_valid:
                print(f"⚠️ Diagram validation failed for '{heading}': {error_msg}")

                # Try to clean up the diagram for A4 compatibility
                print(f"🔧 Attempting to clean up diagram for A4 compatibility...")
                cleaned_mermaid_code = cleanup_diagram_for_a4(fixed_mermaid_code)

                # Re-validate the cleaned diagram
                is_valid_cleaned, error_msg_cleaned, complexity_cleaned = validate_diagram_for_a4(cleaned_mermaid_code, heading)
                if is_valid_cleaned:
                    print(f"✅ Diagram cleaned successfully for A4 compatibility")
                    fixed_mermaid_code = cleaned_mermaid_code
                    complexity = complexity_cleaned
                else:
                    print(f"❌ Diagram still not A4 compatible after cleanup: {error_msg_cleaned}")
                    return None
            else:
                print(f"✅ Diagram passed A4 compatibility validation")

            print(f"📊 Diagram complexity: {complexity['node_count']} nodes, {complexity['subgraph_count']} subgraphs")

            # Only generate sub-diagrams if the original is actually too complex
            if complexity['is_too_complex']:
                print(f"⚠️ Diagram too complex for A4 page, skipping for performance...")
                print(f"📋 Issues: {', '.join(complexity['suggestions'])}")

                # Skip sub-diagram generation for faster processing
                # sub_diagrams = generate_sub_diagrams(heading, content, diagram_type, model, fixed_mermaid_code)
                sub_diagrams = []

                if sub_diagrams and len(sub_diagrams) > 1:
                    print(f"✅ Generated {len(sub_diagrams)} sub-diagrams for better A4 fit")

                    # Store all sub-diagrams in the global list
                    global current_document_diagrams
                    for sub_diagram in sub_diagrams:
                        current_document_diagrams.append({
                            'id': f"{heading.replace(' ', '_')}_{sub_diagram['part_number']}_{len(current_document_diagrams)}",
                            'sectionTitle': sub_diagram['title'],
                            'diagramType': diagram_type,
                            'mermaidCode': sub_diagram['mermaid_code'],
                            'theme': 'default',
                            'lastModified': datetime.now().isoformat(),
                            'isSubDiagram': True,
                            'parentSection': heading,
                            'focusArea': sub_diagram['focus_area']
                        })

                    # Use the first sub-diagram as the main one
                    fixed_mermaid_code = sub_diagrams[0]['mermaid_code']
                    print(f"📊 Using sub-diagram 1 as main: {sub_diagrams[0]['title']}")
                else:
                    print(f"⚠️ Could not break into sub-diagrams, using original")
            else:
                print(f"✅ Diagram is A4-compatible ({complexity['node_count']} nodes) - no sub-diagrams needed")

            # ADD DIAGRAM SIGNATURE TO PREVENT FUTURE DUPLICATES
            diagram_signature = generate_diagram_signature(fixed_mermaid_code, heading)
            generated_diagram_signatures.add(diagram_signature)
            print(f"📝 Added diagram signature: {diagram_signature}")

            # Increment global diagram counter
            total_diagrams_generated += 1

            print(f"✅ Generated Mermaid diagram {total_diagrams_generated}/3 for: {heading}")
            print(f"📋 Original code preview: {mermaid_code[:100]}...")
            print(f"📋 Fixed code preview: {fixed_mermaid_code[:100]}...")
            return fixed_mermaid_code
        else:
            print(f"⚠️ Empty diagram response from Gemini for '{heading}'")
            return None

    except Exception as e:
        print(f"❌ Error generating Mermaid diagram for {heading}: {str(e)}")
        print(f"🔍 Error type: {type(e).__name__}")

        # Provide more specific error messages
        if "API_KEY_INVALID" in str(e):
            print(f"💡 Invalid Gemini API key for diagram generation")
        elif "PERMISSION_DENIED" in str(e):
            print(f"💡 Permission denied for Gemini API")
        elif "QUOTA_EXCEEDED" in str(e):
            print(f"💡 Gemini API quota exceeded")

        return None

def find_mmdc_command():
    """Find the mmdc command path on different systems"""
    possible_commands = [
        'mmdc',  # Direct command
        'mmdc.cmd',  # Windows batch file
        'npx mmdc',  # Using npx
    ]

    # Also check common npm global paths
    import platform
    if platform.system() == "Windows":
        # Common Windows npm global paths
        possible_paths = [
            os.path.expanduser("~\\AppData\\Roaming\\npm\\mmdc.cmd"),
            os.path.expanduser("~\\AppData\\Roaming\\npm\\mmdc"),
            "C:\\Users\\%USERNAME%\\AppData\\Roaming\\npm\\mmdc.cmd",
        ]
        possible_commands.extend(possible_paths)

    for cmd in possible_commands:
        try:
            if cmd.startswith('npx'):
                # For npx commands, split them
                cmd_parts = cmd.split()
                result = subprocess.run(cmd_parts + ['--version'], capture_output=True, text=True, timeout=10)
            else:
                result = subprocess.run([cmd, '--version'], capture_output=True, text=True, timeout=10)

            if result.returncode == 0:
                print(f"✅ Found mmdc at: {cmd}")
                return cmd
        except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
            continue

    return None

def check_mermaid_cli_installed() -> bool:
    """Check if mermaid-cli (mmdc) is installed"""
    return find_mmdc_command() is not None

def diagnose_mermaid_cli_issues():
    """
    Diagnose common Mermaid CLI installation and configuration issues
    """
    print("🔍 Diagnosing Mermaid CLI issues...")

    # Check if Node.js is installed
    try:
        result = subprocess.run(['node', '--version'], capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            print(f"✅ Node.js is installed: {result.stdout.strip()}")
        else:
            print("❌ Node.js is not installed or not working")
            print("💡 Install Node.js from: https://nodejs.org/")
            return
    except Exception:
        print("❌ Node.js is not installed or not accessible")
        print("💡 Install Node.js from: https://nodejs.org/")
        return

    # Check if npm is working
    try:
        result = subprocess.run(['npm', '--version'], capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            print(f"✅ npm is working: {result.stdout.strip()}")
        else:
            print("❌ npm is not working properly")
            return
    except Exception:
        print("❌ npm is not accessible")
        return

    # Check if mmdc is installed
    mmdc_cmd = find_mmdc_command()
    if mmdc_cmd:
        print(f"✅ Found mmdc at: {mmdc_cmd}")

        # Test mmdc with a simple diagram
        test_mermaid = "graph TD\n    A[Start] --> B[End]"
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_file:
                temp_file.write(test_mermaid)
                temp_mmd_path = temp_file.name

            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_png:
                temp_png_path = temp_png.name

            if mmdc_cmd.startswith('npx'):
                cmd = mmdc_cmd.split() + ['-i', temp_mmd_path, '-o', temp_png_path]
            else:
                cmd = [mmdc_cmd, '-i', temp_mmd_path, '-o', temp_png_path]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            # Cleanup
            if os.path.exists(temp_mmd_path):
                os.unlink(temp_mmd_path)
            if os.path.exists(temp_png_path):
                os.unlink(temp_png_path)

            if result.returncode == 0:
                print("✅ mmdc is working correctly")
            else:
                print(f"❌ mmdc test failed: {result.stderr}")
                print("💡 Try reinstalling: npm install -g @mermaid-js/mermaid-cli")

        except Exception as e:
            print(f"❌ mmdc test error: {e}")
    else:
        print("❌ mmdc (Mermaid CLI) is not installed")
        print("💡 Install with: npm install -g @mermaid-js/mermaid-cli")
        print("💡 On Windows, you may need to restart your terminal after installation")
        print("💡 Alternative: Use npx @mermaid-js/mermaid-cli instead of global installation")

def convert_mermaid_to_svg(mermaid_code: str, output_path: str) -> bool:
    """
    Convert Mermaid code to SVG image using mermaid-cli (mmdc)
    SVG is vector-based and stays sharp at any zoom level
    Returns True if successful, False otherwise
    """
    try:
        # Find the mmdc command
        mmdc_cmd = find_mmdc_command()
        if not mmdc_cmd:
            print("❌ mermaid-cli (mmdc) not found.")
            print("💡 To install: npm install -g @mermaid-js/mermaid-cli")
            print("💡 Alternative: Install Node.js first, then run the npm command")
            print("💡 On Windows, you might need to restart your terminal after installation")
            return False

        # Create temporary file for mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_file:
            temp_file.write(mermaid_code)
            temp_mmd_path = temp_file.name

        print(f"🔧 Using mmdc command: {mmdc_cmd}")
        print(f"📝 Temp mermaid file: {temp_mmd_path}")
        print(f"🎯 Output path: {output_path}")

        # Build command based on whether it's npx or direct
        # For SVG, we don't need width/height since it's vector-based
        if mmdc_cmd.startswith('npx'):
            cmd = mmdc_cmd.split() + [
                '-i', temp_mmd_path,
                '-o', output_path,
                '-b', 'white',  # Background color
            ]
        else:
            cmd = [
                mmdc_cmd,
                '-i', temp_mmd_path,
                '-o', output_path,
                '-b', 'white',  # Background color
            ]

        print(f"🚀 Running command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        # Clean up temporary file
        if temp_mmd_path and os.path.exists(temp_mmd_path):
            os.unlink(temp_mmd_path)

        if result.returncode == 0:
            print(f"✅ Successfully converted Mermaid to SVG: {output_path}")
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"📊 Generated SVG size: {file_size} bytes")
                return True
            else:
                print(f"❌ Output file was not created: {output_path}")
                return False
        else:
            print(f"❌ Error converting Mermaid to SVG (exit code: {result.returncode}):")
            print(f"   stdout: {result.stdout}")
            print(f"   stderr: {result.stderr}")
            return False

    except subprocess.TimeoutExpired:
        print("❌ Timeout while converting Mermaid diagram (30s limit exceeded)")
        # Clean up temp file if it exists
        try:
            os.unlink(temp_mmd_path)
        except:
            pass
        return False
    except Exception as e:
        print(f"❌ Error converting Mermaid to SVG: {str(e)}")
        return False

def convert_mermaid_to_png(mermaid_code: str, output_path: str) -> bool:
    """
    Convert Mermaid code to PNG image using mermaid-cli (mmdc)
    This is for document insertion since python-docx doesn't support SVG
    Returns True if successful, False otherwise
    """
    temp_mmd_path = None
    try:
        print(f"🔧 Converting Mermaid to PNG: {output_path}")

        # Find the mmdc command
        mmdc_cmd = find_mmdc_command()
        if not mmdc_cmd:
            print("❌ mmdc command not found")
            return False

        print(f"✅ Using mmdc command: {mmdc_cmd}")

        # Validate Mermaid code first
        if not mermaid_code or not mermaid_code.strip():
            print("❌ Empty Mermaid code provided")
            return False

        # Create temporary file for mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as temp_file:
            temp_file.write(mermaid_code)
            temp_mmd_path = temp_file.name

        print(f"📝 Temp Mermaid file: {temp_mmd_path}")
        print(f"📋 Mermaid code preview: {mermaid_code[:100]}...")

        # Build command for PNG with better error handling
        if mmdc_cmd.startswith('npx'):
            cmd = mmdc_cmd.split() + [
                '-i', temp_mmd_path,
                '-o', output_path,
                '-b', 'white',
            ]
        else:
            cmd = [
                mmdc_cmd,
                '-i', temp_mmd_path,
                '-o', output_path,
                '-b', 'white',
            ]

        print(f"🚀 Running command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=90)

        print(f"📊 Command exit code: {result.returncode}")
        if result.stdout:
            print(f"📤 stdout: {result.stdout}")
        if result.stderr:
            print(f"📥 stderr: {result.stderr}")

        # Clean up temporary file
        if temp_mmd_path and os.path.exists(temp_mmd_path):
            os.unlink(temp_mmd_path)

        if result.returncode == 0:
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"✅ Successfully generated PNG: {output_path} ({file_size} bytes)")
                return True
            else:
                print(f"❌ Command succeeded but output file not found: {output_path}")
                return False
        else:
            print(f"❌ mmdc command failed with exit code {result.returncode}")
            if result.stderr:
                print(f"   Error details: {result.stderr}")
            return False

    except subprocess.TimeoutExpired:
        print("❌ Timeout while converting Mermaid diagram (90s limit exceeded)")
        return False
    except Exception as e:
        print(f"❌ Exception during Mermaid to PNG conversion: {str(e)}")
        print(f"   Exception type: {type(e).__name__}")
        return False
    finally:
        # Clean up temp file if it exists
        if temp_mmd_path and os.path.exists(temp_mmd_path):
            try:
                os.unlink(temp_mmd_path)
            except Exception as cleanup_error:
                print(f"⚠️ Failed to cleanup temp file: {cleanup_error}")

def convert_mermaid_to_png_with_fallback(mermaid_code: str, output_path: str, heading_text: str) -> bool:
    """
    Convert Mermaid to PNG with multiple fallback strategies including Gemini self-correction
    """
    print(f"🔄 Attempting PNG conversion for: {heading_text}")

    # Strategy 1: Try with current settings
    success, error_message = convert_mermaid_to_png_with_error_capture(mermaid_code, output_path)
    if success:
        return True

    print(f"⚠️ First attempt failed: {error_message}")
    print(f"🔧 Trying Gemini self-correction...")

    # Strategy 2: Ask Gemini to fix its own syntax errors
    if error_message:
        corrected_code = fix_mermaid_with_gemini(mermaid_code, error_message)
        if corrected_code != mermaid_code:
            print(f"🔄 Trying with Gemini-corrected code...")
            success, second_error = convert_mermaid_to_png_with_error_capture(corrected_code, output_path)
            if success:
                print(f"✅ Gemini self-correction worked!")
                return True

            # Strategy 2b: If Gemini correction still fails, try one more time with the new error
            if second_error and second_error != error_message:
                print(f"🔧 Gemini correction still failed, trying second correction...")
                double_corrected_code = fix_mermaid_with_gemini(corrected_code, second_error)
                if double_corrected_code != corrected_code:
                    success, _ = convert_mermaid_to_png_with_error_capture(double_corrected_code, output_path)
                    if success:
                        print(f"✅ Gemini double-correction worked!")
                        return True

    print(f"⚠️ Gemini correction failed, trying other fallback strategies...")

    # Strategy 3: Try with simplified command (no extra parameters)
    temp_mmd_path = None
    try:
        mmdc_cmd = find_mmdc_command()
        if not mmdc_cmd:
            return False

        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as temp_file:
            temp_file.write(mermaid_code)
            temp_mmd_path = temp_file.name

        # Simple command without extra parameters
        if mmdc_cmd.startswith('npx'):
            cmd = mmdc_cmd.split() + ['-i', temp_mmd_path, '-o', output_path]
        else:
            cmd = [mmdc_cmd, '-i', temp_mmd_path, '-o', output_path]

        print(f"🔄 Fallback attempt with simplified command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        if temp_mmd_path and os.path.exists(temp_mmd_path):
            os.unlink(temp_mmd_path)

        if result.returncode == 0 and os.path.exists(output_path):
            print(f"✅ Fallback conversion successful!")
            return True

    except Exception as e:
        print(f"❌ Fallback strategy failed: {e}")
        if temp_mmd_path and os.path.exists(temp_mmd_path):
            try:
                os.unlink(temp_mmd_path)
            except:
                pass

    # Strategy 4: Try to fix common Mermaid syntax issues
    print(f"🔄 Trying syntax cleanup...")
    cleaned_code = clean_mermaid_syntax(mermaid_code)
    if cleaned_code != mermaid_code:
        print(f"📝 Cleaned Mermaid syntax, retrying...")
        if convert_mermaid_to_png(cleaned_code, output_path):
            return True

    print(f"❌ All fallback strategies failed for: {heading_text}")
    return False

def clean_mermaid_syntax(mermaid_code: str) -> str:
    """
    Clean common Mermaid syntax issues that might cause failures
    """
    # Just call the main fix function to avoid conflicts
    return fix_mermaid_syntax(mermaid_code)

def insert_diagram_placeholder(doc: Document, heading_text: str, mermaid_code: str):
    """
    Insert a placeholder when diagram generation fails
    """
    try:
        # Add a note about the diagram failure
        placeholder_para = doc.add_paragraph()
        placeholder_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder_para.paragraph_format.space_before = Pt(12)
        placeholder_para.paragraph_format.space_after = Pt(12)

        # Add placeholder text
        run = placeholder_para.add_run("📊 Diagram Generation Failed")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

        # Add explanation
        explanation_para = doc.add_paragraph(
            f"A diagram was intended for this section but could not be generated. "
            f"You can use the Draw.io XML export feature to create the diagram manually.",
            style='SRS Normal'
        )
        explanation_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        explanation_para.paragraph_format.space_after = Pt(12)

        # Add the Mermaid code for reference
        code_para = doc.add_paragraph("Mermaid Code Reference:", style='SRS Normal')
        code_para.paragraph_format.space_before = Pt(6)

        code_content = doc.add_paragraph(mermaid_code[:500] + "..." if len(mermaid_code) > 500 else mermaid_code)
        code_content.paragraph_format.left_indent = Inches(0.5)
        code_content.paragraph_format.space_after = Pt(12)

        # Style the code block
        for run in code_content.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        print(f"📝 Inserted diagram placeholder for: {heading_text}")

    except Exception as e:
        print(f"❌ Failed to insert diagram placeholder: {e}")

def insert_diagram_into_document(doc: Document, diagram_path: str, heading_text: str):
    """
    Insert diagram image into the document after the specified heading
    Ensures A4 compatibility with proper sizing
    """
    try:
        if not os.path.exists(diagram_path):
            print(f"❌ Diagram file not found: {diagram_path}")
            return

        # Find the heading paragraph and insert diagram after content
        for i, para in enumerate(doc.paragraphs):
            if heading_text.lower() in para.text.lower() and para.style.name == 'SRS Heading 2':
                # Add diagram description
                diagram_desc = doc.add_paragraph("The following diagram illustrates the system design:", style='SRS Normal')
                diagram_desc.paragraph_format.left_indent = Inches(0.25)
                diagram_desc.paragraph_format.space_after = Pt(6)

                # Add the diagram image with A4-optimized sizing
                diagram_para = doc.add_paragraph()
                diagram_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = diagram_para.add_run()

                # A4 OPTIMIZATION: Calculate optimal width for A4 page
                # A4 width: 8.27 inches, margins: ~1 inch each side, so usable width: ~6.27 inches
                # Use 6 inches for diagrams to ensure they fit comfortably
                optimal_width = Inches(6.0)

                run.add_picture(diagram_path, width=optimal_width)
                diagram_para.paragraph_format.space_after = Pt(12)

                print(f"✅ Inserted A4-optimized diagram for: {heading_text}")
                break

    except Exception as e:
        print(f"❌ Error inserting diagram for {heading_text}: {str(e)}")

def insert_diagram_with_a4_optimization(doc: Document, diagram_path: str, heading_text: str, diagram_type: str = "flowchart"):
    """
    Insert diagram with A4 page optimization and automatic sizing
    """
    try:
        if not os.path.exists(diagram_path):
            print(f"❌ Diagram file not found: {diagram_path}")
            return False

        # Find the heading paragraph
        target_paragraph = None
        for para in doc.paragraphs:
            if heading_text.lower() in para.text.lower() and para.style.name == 'SRS Heading 2':
                target_paragraph = para
                break

        if not target_paragraph:
            print(f"⚠️ Could not find section '{heading_text}' in document")
            return False

        # Add diagram description
        diagram_desc = doc.add_paragraph("The following diagram illustrates the system design:", style='SRS Normal')
        diagram_desc.paragraph_format.left_indent = Inches(0.25)
        diagram_desc.paragraph_format.space_after = Pt(6)

        # A4 OPTIMIZATION: Smart sizing based on diagram type
        if diagram_type == "erDiagram":
            # ER diagrams are typically wider, use smaller width
            diagram_width = Inches(5.5)
        elif diagram_type == "sequenceDiagram":
            # Sequence diagrams can be wider, use medium width
            diagram_width = Inches(6.0)
        else:
            # Flowcharts and others use standard A4 width
            diagram_width = Inches(6.0)

        # Insert the diagram with optimized sizing
        diagram_para = doc.add_paragraph()
        diagram_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = diagram_para.add_run()

        try:
            run.add_picture(diagram_path, width=diagram_width)
            diagram_para.paragraph_format.space_after = Pt(12)
            print(f"✅ Successfully inserted A4-optimized {diagram_type} diagram into document")
            print(f"📏 Diagram width: {diagram_width}")
            return True

        except Exception as png_error:
            print(f"❌ Failed to insert diagram: {png_error}")
            return False

    except Exception as e:
        print(f"❌ Error in A4-optimized diagram insertion: {e}")
        return False

def generate_srs_docx(headings: List[Dict[str, Any]], output_path: str, uploaded_content: str = "", project_title: str = ""):
    """
    Generate SRS document in DOCX format from selected headings

    Args:
        headings: List of heading dictionaries with 'heading', 'purpose', 'category', 'source'
        output_path: Path where to save the generated DOCX file
        uploaded_content: Content from uploaded PDF documents for context

    Returns:
        List of generated diagrams with their metadata
    """
    global current_document_diagrams

    # Reset diagram generation for new document
    reset_diagram_counter()
    current_document_diagrams = []  # Reset for new document

    try:
        # Create a new document
        doc = Document()

        # Set up document styles
        setup_document_styles(doc)

        # Set up header and footer
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'srs_logo.png')
        setup_header_footer(doc, logo_path)

        # Add title page with logo and project name
        add_title_page(doc, project_title)

        # Add table of contents
        add_table_of_contents(doc, headings)

        # Add content sections
        add_content_sections(doc, headings, uploaded_content)

        # Save the document
        doc.save(output_path)

        print(f"✅ SRS document generated successfully: {output_path}")

        # Force footer refresh to ensure changes take effect
        force_footer_refresh(doc)

        # Print diagram generation summary
        print_diagram_generation_summary()

        # Return the generated diagrams
        return current_document_diagrams

    except Exception as e:
        print(f"❌ Failed to generate SRS document: {str(e)}")
        raise

def setup_document_styles(doc: Document):
    """Set up document styles for consistent formatting"""
    try:
        # Title style
        title_style = doc.styles.add_style('SRS Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Cambria'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)  # Black
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.line_spacing = 1.0  # Single spacing

        # Heading 1 style - Light Blue
        h1_style = doc.styles.add_style('SRS Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Cambria'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(70, 130, 180)  # Light Blue (Steel Blue)
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        h1_style.paragraph_format.line_spacing = 1.0  # Single spacing

        # Heading 2 style - Light Blue
        h2_style = doc.styles.add_style('SRS Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Cambria'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(70, 130, 180)  # Light Blue (Steel Blue)
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(6)
        h2_style.paragraph_format.line_spacing = 1.0  # Single spacing

        # Heading 3 style - Light Blue
        h3_style = doc.styles.add_style('SRS Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Cambria'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(70, 130, 180)  # Light Blue (Steel Blue)
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)
        h3_style.paragraph_format.line_spacing = 1.0  # Single spacing

        # Normal text style - Black
        normal_style = doc.styles.add_style('SRS Normal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Cambria'
        normal_font.size = Pt(11)
        normal_font.color.rgb = RGBColor(0, 0, 0)  # Black
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing = 1.0  # Single spacing

    except Exception as e:
        print(f"⚠️ Warning: Could not set up custom styles: {e}")

def setup_header_footer(doc: Document, logo_path: str = None):
    """Set up document header and footer for ALL sections"""
    try:
        # Apply to ALL sections, not just the first one
        for section_idx, section in enumerate(doc.sections):
            print(f"🔧 Setting up header/footer for section {section_idx + 1}/{len(doc.sections)}")

            # Setup Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right align for logo

            if logo_path and os.path.exists(logo_path):
                # Add logo to header (right side)
                run = header_para.runs[0] if header_para.runs else header_para.add_run()
                run.add_picture(logo_path, width=Inches(2.0))
            else:
                # Add text logo if no image available (right aligned)
                header_para.text = "SRS Document"
                for run in header_para.runs:
                    run.font.bold = True
                    run.font.size = Pt(14)

            # Ensure footers appear on all pages (no special first-page footer)
            try:
                section.different_first_page_header_footer = False
            except Exception:
                pass
            try:
                section.footer_distance = Inches(0.4)
            except Exception:
                pass

            # Compute content width (page width minus margins)
            page_width = section.page_width.inches
            left_margin = section.left_margin.inches
            right_margin = section.right_margin.inches
            content_width = page_width - left_margin - right_margin

            def _build_footer(ftr):
                if not ftr:
                    return

                # Clear footer completely (remove all existing paragraphs)
                for para in ftr.paragraphs:
                    p = para._element
                    p.getparent().remove(p)

                # Create one clean paragraph
                p = ftr.add_paragraph()
                p.style.font.size = Pt(9)

                # PERFECT FOOTER PARAMETERS - analyzed from your example
                # A4 page: 21cm width, typical margins: 2.5cm left + 2.5cm right = 16cm content
                # But footer should use FULL page width for maximum spread

                tabs = p.paragraph_format.tab_stops
                tabs.clear_all()
                from docx.shared import Cm

                # Set paragraph formatting for maximum spread
                pf = p.paragraph_format
                pf.left_indent = Cm(-2.0)   # Push far beyond normal left margin
                pf.right_indent = Cm(-2.0)  # Pull far beyond normal right margin
                pf.first_line_indent = Cm(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)

                # Right tab positioned slightly more left (17cm for better A4 visibility)
                tabs.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)

                # Add left text
                left_run = p.add_run("Software Requirement Specification")
                left_run.font.name = 'Cambria'
                left_run.font.size = Pt(10)
                left_run.font.color.rgb = RGBColor(0, 0, 0)

                # Add tab (moves cursor to right tab stop)
                p.add_run("\t")

                # Add right text with page fields (7pt dark format)
                page_run = p.add_run("Page ")
                page_run.font.name = 'Cambria'
                page_run.font.size = Pt(7)  # 7pt font size
                page_run.font.color.rgb = RGBColor(0, 0, 0)  # Dark color

                # Helper function to add fields (your exact implementation)
                def add_field(run, field_type):
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText')
                    instrText.text = field_type
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'separate')
                    fldChar3 = OxmlElement('w:fldChar')
                    fldChar3.set(qn('w:fldCharType'), 'end')
                    run._r.append(fldChar1)
                    run._r.append(instrText)
                    run._r.append(fldChar2)
                    run._r.append(fldChar3)

                # Add PAGE field (7pt dark format)
                r1 = p.add_run()
                r1.font.name = 'Cambria'
                r1.font.size = Pt(7)  # 7pt font size
                r1.font.color.rgb = RGBColor(0, 0, 0)  # Dark color
                add_field(r1, "PAGE")

                # Add " of " (7pt dark format)
                of_run = p.add_run(" of ")
                of_run.font.name = 'Cambria'
                of_run.font.size = Pt(7)  # 7pt font size
                of_run.font.color.rgb = RGBColor(0, 0, 0)  # Dark color

                # Add NUMPAGES field (7pt dark format)
                r2 = p.add_run()
                r2.font.name = 'Cambria'
                r2.font.size = Pt(7)  # 7pt font size
                r2.font.color.rgb = RGBColor(0, 0, 0)  # Dark color
                add_field(r2, "NUMPAGES")

            # Build footer for all variants to be safe
            _build_footer(getattr(section, 'footer', None))
            _build_footer(getattr(section, 'first_page_footer', None))
            _build_footer(getattr(section, 'even_page_footer', None))

        print("✅ Header and footer setup completed for all sections")

    except Exception as e:
        print(f"⚠️ Warning: Could not set up header/footer: {e}")

def force_footer_refresh(doc: Document):
    """Force refresh of footers to ensure changes take effect"""
    try:
        print("🔄 Forcing footer refresh across all sections...")
        for section_idx, section in enumerate(doc.sections):
            # Force footer to be different first page = False
            section.different_first_page_header_footer = False

            # Ensure footer distance is set
            section.footer_distance = Inches(0.4)

            print(f"✅ Section {section_idx + 1} footer refreshed")

        print("✅ All footers refreshed successfully")
    except Exception as e:
        print(f"⚠️ Warning: Could not refresh footers: {e}")

def _add_markdown_runs(paragraph, text: str):
    """Render enhanced Markdown (**bold**, *italic*) into python-docx runs.
    If no markdown tokens present, add plain text run.
    """
    import re
    tokens = []
    i = 0
    while i < len(text):
        if text.startswith('**', i):
            j = text.find('**', i+2)
            if j != -1:
                tokens.append(('bold', text[i+2:j]))
                i = j + 2
                continue
        if text.startswith('*', i) and not text.startswith('**', i):
            j = text.find('*', i+1)
            if j != -1 and not text.startswith('**', j-1):
                tokens.append(('italic', text[i+1:j]))
                i = j + 1
                continue
        # plain text chunk
        next_special = len(text)
        n1 = text.find('**', i)
        n2 = text.find('*', i)
        candidates = [n for n in [n1, n2] if n != -1]
        if candidates:
            next_special = min(candidates)
        tokens.append(('plain', text[i:next_special]))
        i = next_special

    # write runs
    if not tokens:
        run = paragraph.add_run(text)
        # Apply font formatting based on paragraph style
        _apply_font_formatting(run, paragraph)
        return

    for token_type, token_text in tokens:
        if not token_text:
            continue
        run = paragraph.add_run(token_text)

        # Apply base font formatting
        _apply_font_formatting(run, paragraph)

        # Apply markdown formatting
        if token_type == 'bold':
            run.font.bold = True
        elif token_type == 'italic':
            run.font.italic = True

def _apply_font_formatting(run, paragraph):
    """Apply consistent font formatting based on paragraph style"""
    style_name = paragraph.style.name if paragraph.style else 'Normal'

    # Set font name to Cambria
    run.font.name = 'Cambria'

    # Set colors based on style
    if 'Heading' in style_name:
        # Light blue for headings
        run.font.color.rgb = RGBColor(70, 130, 180)  # Steel Blue
    else:
        # Black for content
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black

def _clean_metadata_from_content(content: str) -> str:
    """Remove any metadata labels and unwanted numbering that might appear in AI-generated content"""
    if not content:
        return content

    lines = content.split('\n')
    cleaned_lines = []

    import re

    for line in lines:
        original_line = line
        line = line.strip()

        # Skip lines that start with metadata labels
        if (line.startswith('Purpose:') or
            line.startswith('Source:') or
            line.startswith('SECTION HEADING:') or
            line.startswith('Section Heading:') or
            line.startswith('Heading:')):
            continue

        # Remove X.Y numbering at the start of lines (1.1, 1.2, etc.)
        line = re.sub(r'^\d+\.\d+\s+', '', line)

        # If line is not empty after cleaning, add it
        if line.strip():
            cleaned_lines.append(line)
        elif original_line.strip() == '':
            # Preserve empty lines for formatting
            cleaned_lines.append('')

    return '\n'.join(cleaned_lines)

def _is_content_subheading(line: str) -> bool:
    """Detect if a line is a sub-heading within content that should be bold"""
    line = line.strip()

    # Skip empty lines and bullet points
    if not line or line.startswith('•') or line.startswith('-'):
        return False

    # Check for common sub-heading patterns
    import re

    # Pattern 1: Lines that end with colon (e.g., "System Requirements:")
    if line.endswith(':') and len(line) > 5 and len(line) < 100:
        return True

    # Pattern 2: Lines that are short and look like titles (no periods, proper case)
    if (len(line) < 80 and
        not line.endswith('.') and
        not line.startswith('The ') and
        not line.startswith('This ') and
        any(word[0].isupper() for word in line.split() if len(word) > 2)):

        # Check if it contains typical sub-heading words
        subheading_words = ['Requirements', 'Specifications', 'Overview', 'Details',
                           'Features', 'Components', 'Architecture', 'Design',
                           'Implementation', 'Configuration', 'Management']

        if any(word in line for word in subheading_words):
            return True

    # Pattern 3: Lines that start with capital letters and are relatively short
    if (line[0].isupper() and
        len(line) < 60 and
        not line.endswith('.') and
        ' shall ' not in line.lower() and
        not line.startswith('The system')):
        return True

    return False
def _add_markdown_block(doc: Document, text: str, style_name: str = 'SRS Normal', justify_text: bool = True):
    """Add a multi-line text block with enhanced markdown support and justification.
    Supports **bold**, *italic*, bullets, markdown headings, and text justification.
    Sub-headings within content are made bold.
    """
    lines = text.splitlines() if text else []
    if not lines:
        doc.add_paragraph('', style=style_name)
        return

    import re
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip('\r')

        # Blank line -> blank paragraph
        if not line.strip():
            doc.add_paragraph('', style=style_name)
            i += 1
            continue

        stripped = line.strip()

        # Detect Markdown table block starting at this line
        def _is_separator_row(row: str) -> bool:
            row = row.strip()
            if not (row.startswith('|') and row.endswith('|')):
                return False
            parts = [c.strip() for c in row.strip('|').split('|')]
            if not parts:
                return False
            for p in parts:
                # ---, :---, ---:, :---:
                if not re.fullmatch(r':?-{3,}:?', p):
                    return False
            return True

        if stripped.startswith('|') and stripped.endswith('|') and (i + 1) < len(lines) and _is_separator_row(lines[i + 1]):
            # Collect the table lines until a non-pipe or empty line
            table_lines = [lines[i].strip(), lines[i + 1].strip()]
            j = i + 2
            while j < len(lines):
                ln = lines[j].strip()
                if ln and ln.startswith('|') and '|' in ln:
                    table_lines.append(ln)
                    j += 1
                else:
                    break

            # Parse header and rows
            header_cells = [c.strip() for c in table_lines[0].strip('|').split('|')]
            data_rows = []
            for r in table_lines[2:]:  # skip separator row
                cells = [c.strip() for c in r.strip('|').split('|')]
                # pad/truncate to header count
                if len(cells) < len(header_cells):
                    cells += [''] * (len(header_cells) - len(cells))
                elif len(cells) > len(header_cells):
                    cells = cells[:len(header_cells)]
                data_rows.append(cells)

            # Build docx table
            rows_count = 1 + len(data_rows)
            cols_count = max(1, len(header_cells))
            tbl = doc.add_table(rows=rows_count, cols=cols_count)
            try:
                tbl.style = 'Table Grid'
            except Exception:
                pass
            tbl.autofit = True

            # Fill header row (bold)
            for col, text_cell in enumerate(header_cells):
                cell = tbl.cell(0, col)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(text_cell)
                run.font.name = 'Cambria'
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Fill data rows
            for r_idx, row_vals in enumerate(data_rows, start=1):
                for c_idx, text_cell in enumerate(row_vals):
                    cell = tbl.cell(r_idx, c_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(text_cell)
                    run.font.name = 'Cambria'
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # Advance index past the table block
            i = j
            continue

        # Detect markdown headings
        heading_level = 0
        content = line
        if line.lstrip().startswith('### '):
            heading_level = 3
            content = line.lstrip()[4:]
        elif line.lstrip().startswith('## '):
            heading_level = 2
            content = line.lstrip()[3:]
        elif line.lstrip().startswith('# '):
            heading_level = 1
            content = line.lstrip()[2:]

        # Check if this looks like a sub-heading (even without markdown syntax)
        is_subheading = _is_content_subheading(line)

        # If it's a markdown heading, use appropriate heading style
        if heading_level > 0:
            if heading_level == 1:
                p = doc.add_paragraph(style='SRS Heading 1')
            elif heading_level == 2:
                p = doc.add_paragraph(style='SRS Heading 2')
            else:  # heading_level >= 3
                p = doc.add_paragraph(style='SRS Heading 3')
            _add_markdown_runs(p, content)
            i += 1
            continue

        # If it's a content sub-heading, make it bold
        elif is_subheading:
            p = doc.add_paragraph(style=style_name)
            if justify_text:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Sub-headings should be left-aligned

            # Make the entire sub-heading bold
            run = p.add_run(content)
            _apply_font_formatting(run, p)
            run.font.bold = True
            i += 1
            continue

        # Detect simple bullets
        bullet = False
        if line.lstrip().startswith('- '):
            bullet = True
            content = line.lstrip()[2:]
        elif line.lstrip().startswith('* '):
            bullet = True
            content = line.lstrip()[2:]
        else:
            content = line

        p = doc.add_paragraph(style=style_name)

        # Set text justification for regular paragraphs
        if justify_text and not bullet:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if bullet:
            p.add_run('• ')

        # Add markdown-formatted content
        _add_markdown_runs(p, content)

        i += 1

def process_markdown_to_docx(doc: Document, markdown_text: str, style_name: str = 'Normal'):
    """Convert markdown text to docx format with enhanced formatting.
    Supports headings, bold, italic, lists, and code blocks.
    """
    try:
        # Ensure basic styles exist, fallback to built-in styles if SRS styles don't exist
        try:
            doc.styles['SRS Normal']
            style_name = 'SRS Normal'
            heading_style_1 = 'SRS Heading 1'
            heading_style_2 = 'SRS Heading 2'
        except KeyError:
            # Fallback to built-in styles
            style_name = 'Normal'
            heading_style_1 = 'Heading 1'
            heading_style_2 = 'Heading 2'

        # Convert markdown to HTML
        html_output = markdown.markdown(markdown_text)

        # Parse HTML to extract plain text with formatting
        soup = BeautifulSoup(html_output, 'html.parser')

        # Process each element
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'code', 'pre']):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Handle headings
                level = int(element.name[1])
                heading_style = heading_style_1 if level <= 2 else heading_style_2
                p = doc.add_paragraph(style=heading_style)
                _add_html_formatted_text(p, element)

            elif element.name == 'p':
                # Handle paragraphs
                p = doc.add_paragraph(style=style_name)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _add_html_formatted_text(p, element)

            elif element.name in ['ul', 'ol']:
                # Handle lists
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style=style_name)
                    p.add_run('• ')
                    _add_html_formatted_text(p, li)

            elif element.name in ['code', 'pre']:
                # Handle code blocks
                p = doc.add_paragraph(style=style_name)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

    except Exception as e:
        print(f"⚠️ Warning: Could not process markdown: {e}")
        # Fallback to simple text processing
        _add_markdown_block(doc, markdown_text, style_name)

def _add_html_formatted_text(paragraph, element):
    """Add HTML formatted text to a paragraph with proper formatting."""
    for content in element.contents:
        if hasattr(content, 'name'):
            if content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.font.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.font.italic = True
            elif content.name == 'code':
                run = paragraph.add_run(content.get_text())
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(content.get_text())
        else:
            # Plain text
            paragraph.add_run(str(content))

def add_title_page(doc: Document, project_title: str = ""):
    """Add title page to the document with logo and project name"""

    # Add logo at the top center
    logo_path = os.path.join(os.path.dirname(__file__), '..', 'srs_logo.png')
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(3.0))  # Larger logo for title page

        # Add spacing after logo
        doc.add_paragraph()

    # Add main title
    doc.add_paragraph('Software Requirements Specification', style='SRS Title')

    # Add spacing
    doc.add_paragraph()

    # Add project name prominently (if provided)
    if project_title:
        proj_para = doc.add_paragraph(style='SRS Title')  # Use title style for prominence
        proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        proj_run = proj_para.add_run(project_title)
        proj_run.font.size = Pt(20)  # Slightly smaller than main title
        proj_run.font.bold = True
        proj_run.font.color.rgb = RGBColor(70, 130, 180)  # Light blue like headings

        # Add spacing after project name
        doc.add_paragraph()

    # Add subtitle
    subtitle = doc.add_paragraph('Generated by SRS Dynamic Generator', style='SRS Normal')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add more spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Add document info
    info_para = doc.add_paragraph('Document Information:', style='SRS Normal')
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Make it bold to stand out
    for run in info_para.runs:
        run.font.bold = True

    # Add document details
    details = [
        ('Document Type:', 'Software Requirements Specification'),
        ('Tool:', 'SRS Dynamic Generator'),
        ('Version:', '1.0')
    ]

    for label, value in details:
        doc.add_paragraph(f'{label} {value}', style='SRS Normal')

    # Add page break
    doc.add_page_break()

def _strip_leading_numbers_preserve_markdown(text: str) -> str:
    """Remove leading numbering like '2.' or '2.1 - ' from a heading string while preserving **bold** wrappers.
    Examples:
    '**2. Heading**' -> '**Heading**'
    '2.1 Heading' -> 'Heading'
    """
    import re
    if not text:
        return text
    # If bold-wrapped, clean inside
    if text.startswith('**') and text.endswith('**') and len(text) >= 4:
        inner = text[2:-2]
        inner = re.sub(r'^\s*\d+(?:\.\d+)*\s*[-\.)]?\s*', '', inner)
        return f"**{inner}**"
    # General case
    return re.sub(r'^\s*\d+(?:\.\d+)*\s*[-\.)]?\s*', '', text)

def _strip_markdown_tokens(text: str) -> str:
    """Remove simple markdown markers (** and *) from a string."""
    if not text:
        return text
    return text.replace('**', '').replace('*', '')

def add_table_of_contents(doc: Document, headings: List[Dict[str, Any]]):
    """Add table of contents to the document"""
    print(f"🔍 TOC DEBUG: Received {len(headings)} headings:")
    for i, heading in enumerate(headings):
        print(f"   TOC {i+1}. {heading.get('heading', 'No heading')} ({heading.get('category', 'No category')})")

    # DEDUPLICATE HEADINGS - Remove duplicates based on heading name and category
    seen_headings = set()
    deduplicated_headings = []

    for heading in headings:
        heading_key = (heading.get('heading', ''), heading.get('category', ''))
        if heading_key not in seen_headings:
            seen_headings.add(heading_key)
            deduplicated_headings.append(heading)
        else:
            print(f"⚠️ TOC: Skipping duplicate heading: {heading.get('heading', '')} ({heading.get('category', '')})")

    print(f"🔧 TOC: Deduplicated {len(headings)} → {len(deduplicated_headings)} headings")

    # Add TOC title
    toc_title = doc.add_paragraph('Table of Contents', style='SRS Heading 1')
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Group all headings by category, mapping custom sections to appropriate categories
    categories = {}
    for heading in deduplicated_headings:
        category = heading.get('category', 'Other')

        # Map custom sections to appropriate existing categories based on content
        if category == 'Custom':
            heading_text = heading.get('heading', '').lower()
            # Try to categorize custom sections intelligently
            if any(word in heading_text for word in ['requirement', 'functional', 'feature']):
                category = 'Requirements'
            elif any(word in heading_text for word in ['design', 'architecture', 'system']):
                category = 'Design'
            elif any(word in heading_text for word in ['interface', 'ui', 'user']):
                category = 'Interface'
            elif any(word in heading_text for word in ['performance', 'quality', 'non-functional']):
                category = 'Non-Functional Requirements'
            else:
                # Default to Requirements for unclassified custom sections
                category = 'Requirements'

        if category not in categories:
            categories[category] = []
        categories[category].append(heading)

    # Define desired category order for numbering
    desired_order = [
        'Introduction',
        'Overall Description',
        'System Features and Functional Requirements',
        'External Interface Requirements',
        'Non-Functional Requirements',
        'System Architecture and Design',
        'Data Requirements',
        'Constraints and Dependencies',
        'Acceptance and Verification',
        'Appendices'
    ]
    ordered_categories = [c for c in desired_order if c in categories] + [c for c in categories.keys() if c not in desired_order]

    # Add TOC entries for categorized headings with numbering X, X.1, X.2, etc.
    for top_idx, category in enumerate(ordered_categories, 1):
        category_headings = categories[category]
        # Add category header with X (not X.0) - use Normal style for TOC
        toc_category = doc.add_paragraph(f'{top_idx} {category}', style='SRS Normal')
        toc_category.runs[0].font.bold = True

        # Add headings in this category
        for sub_idx, heading in enumerate(category_headings, 1):
            heading_text = heading.get('heading', '')
            cleaned = _strip_markdown_tokens(_strip_leading_numbers_preserve_markdown(heading_text))
            toc_entry = doc.add_paragraph(f'{top_idx}.{sub_idx} {cleaned}', style='SRS Normal')
            toc_entry.paragraph_format.left_indent = Inches(0.5)

    # Custom sections are now included in regular categories, no special handling needed

    # Add page break
    doc.add_page_break()

def add_content_sections(doc: Document, headings: List[Dict[str, Any]], uploaded_content: str = ""):
    """Add content sections for each heading with full SRS context for diagrams"""
    print(f"🔍 DEBUG: Received {len(headings)} headings:")
    for i, heading in enumerate(headings):
        print(f"   {i+1}. {heading.get('heading', 'No heading')} ({heading.get('category', 'No category')})")

    # DEDUPLICATE HEADINGS - Remove duplicates based on heading name and category
    seen_headings = set()
    deduplicated_headings = []

    for heading in headings:
        heading_key = (heading.get('heading', ''), heading.get('category', ''))
        if heading_key not in seen_headings:
            seen_headings.add(heading_key)
            deduplicated_headings.append(heading)
        else:
            print(f"⚠️ CONTENT: Skipping duplicate heading: {heading.get('heading', '')} ({heading.get('category', '')})")

    print(f"🔧 CONTENT: Deduplicated {len(headings)} → {len(deduplicated_headings)} headings")

    # Group all headings by category, mapping custom sections to appropriate categories
    categories = {}
    for heading in deduplicated_headings:
        category = heading.get('category', 'Other')

        # Map custom sections to appropriate existing categories based on content
        if category == 'Custom':
            heading_text = heading.get('heading', '').lower()
            # Try to categorize custom sections intelligently
            if any(word in heading_text for word in ['requirement', 'functional', 'feature']):
                category = 'Requirements'
            elif any(word in heading_text for word in ['design', 'architecture', 'system']):
                category = 'Design'
            elif any(word in heading_text for word in ['interface', 'ui', 'user']):
                category = 'Interface'
            elif any(word in heading_text for word in ['performance', 'quality', 'non-functional']):
                category = 'Non-Functional Requirements'
            else:
                # Default to Requirements for unclassified custom sections
                category = 'Requirements'

        if category not in categories:
            categories[category] = []
        categories[category].append(heading)

    # FIRST PASS: Generate all content without diagrams to build full SRS context
    print("🔄 First pass: Generating all content for full SRS context...")
    all_generated_content = []
    all_headings_for_processing = []

    # Collect all headings in order
    for category, category_headings in categories.items():
        for heading in category_headings:
            all_headings_for_processing.append(heading)

    # Generate content for all headings first (limit to first 5 for performance)
    print(f"📊 Processing {len(all_headings_for_processing)} headings (limiting to first 5 for performance)")
    limited_headings = all_headings_for_processing[:5]  # Limit to first 5 headings

    for heading in limited_headings:
        heading_text = heading.get('heading', '')
        purpose = heading.get('purpose', '')
        source = heading.get('source', 'Unknown')
        user_prompt = heading.get('userPrompt', '')

        print(f"📝 Pre-generating content for: {heading_text}")
        generated_content = generate_content_for_heading(
            heading_text, purpose, source, uploaded_content, user_prompt
        )

        all_generated_content.append({
            'heading': heading_text,
            'content': generated_content,
            'heading_data': heading
        })

    # Build full SRS context string
    full_srs_content = "\n\n".join([
        f"=== {item['heading']} ===\n{item['content']}"
        for item in all_generated_content
    ])

    print(f"📊 Full SRS context built: {len(full_srs_content)} characters")

    # SECOND PASS: Add content to document with full SRS context for diagrams
    print("🔄 Second pass: Adding pre-generated content to document with diagram generation...")

    # Create a lookup for pre-generated content
    content_lookup = {item['heading']: item for item in all_generated_content}

    # Define desired category order for consistent numbering
    desired_order = [
        'Introduction',
        'Overall Description',
        'System Features and Functional Requirements',
        'External Interface Requirements',
        'Non-Functional Requirements',
        'System Architecture and Design',
        'Data Requirements',
        'Constraints and Dependencies',
        'Acceptance and Verification',
        'Appendices'
    ]
    ordered_categories = [c for c in desired_order if c in categories] + [c for c in categories.keys() if c not in desired_order]

    # Add content for each category (excluding custom sections) with numbering X, X.1, X.2, etc.
    for top_idx, category in enumerate(ordered_categories, 1):
        category_headings = categories[category]
        # Add category header with numbering (X not X.0)
        doc.add_paragraph(f'{top_idx} {category}', style='SRS Heading 1')

        # Add content for each heading in this category
        for sub_idx, heading in enumerate(category_headings, 1):
            heading_text = heading.get('heading', '')
            # Inject computed numbering into heading dict for downstream rendering
            heading_with_number = dict(heading)
            heading_with_number['_number'] = f'{top_idx}.{sub_idx}'
            pre_generated = content_lookup.get(heading_text)
            if pre_generated:
                add_heading_content_with_pregenerated(doc, heading_with_number, uploaded_content, full_srs_content, pre_generated['content'])
            else:
                # Fallback to old method if not found
                add_heading_content_with_full_context(doc, heading_with_number, uploaded_content, full_srs_content)

        # Add some spacing between categories
        doc.add_paragraph()

    # Custom sections are now included in regular categories, no special processing needed

def add_heading_content_with_pregenerated(doc: Document, heading: Dict[str, Any], uploaded_content: str = "", full_srs_content: str = "", pregenerated_content: str = ""):
    """Add content for a specific heading with diagram generation using pre-generated content"""
    global current_document_diagrams

    heading_text = heading.get('heading', '')
    purpose = heading.get('purpose', '')
    source = heading.get('source', 'Unknown')
    category = heading.get('category', 'Other')
    user_prompt = heading.get('userPrompt', '')

    # Add heading (supports **bold** and *italic*; apply computed numbering X.Y if provided)
    number = heading.get('_number')
    display_text = _strip_leading_numbers_preserve_markdown(heading_text)
    if number:
        display_text = f"{number} {display_text}"
    heading_para = doc.add_paragraph(style='SRS Heading 2')
    _add_markdown_runs(heading_para, display_text)

    # Remove Purpose and Source metadata as requested
    # Add spacing after heading
    heading_para.paragraph_format.space_after = Pt(12)

    # Use pre-generated content (no need to call AI again)
    print(f"📄 Using pre-generated content for: {heading_text}")
    generated_content = pregenerated_content

    # Clean any remaining metadata from content
    generated_content = _clean_metadata_from_content(generated_content)

    # Add generated content with markdown formatting support
    _add_markdown_block(doc, generated_content, style_name='SRS Normal')

    # Check if this section should have a diagram (considering user prompt)
    diagram_type = should_generate_diagram(heading_text, user_prompt)
    if diagram_type:
        print(f"🎨 Generating diagram for section: {heading_text}")

        # Generate Mermaid diagram with custom prompt support and full SRS context
        mermaid_code = generate_mermaid_diagram(heading_text, generated_content, uploaded_content, user_prompt, full_srs_content)

        # If sequence diagram generation failed, try with a more specific prompt
        if not mermaid_code and diagram_type == 'sequenceDiagram':
            print(f"🔄 Sequence diagram generation failed, trying with specialized prompt...")
            mermaid_code = generate_sequence_diagram_specialized(heading_text, generated_content, uploaded_content, user_prompt, full_srs_content)

        if False and mermaid_code:
            # Disabled legacy convert/insert path to prevent duplicate insertions
            pass
        else:
            print(f"⚠️ No Mermaid code generated for: {heading_text}")

        if mermaid_code:
            # Store the generated diagram
            current_document_diagrams.append({
                'id': f"{heading_text.replace(' ', '_')}_{len(current_document_diagrams)}",
                'sectionTitle': heading_text,
                'diagramType': diagram_type,
                'mermaidCode': mermaid_code,
                'theme': 'default',
                'lastModified': datetime.now().isoformat()
            })
            # Create temporary directory for diagrams if it doesn't exist
            diagrams_dir = os.path.join(os.path.dirname(__file__), '..', 'temp_diagrams')
            os.makedirs(diagrams_dir, exist_ok=True)

            # Generate unique filename for the diagram
            safe_heading = "".join(c for c in heading_text if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_heading = safe_heading.replace(' ', '_')
            diagram_filename = f"{safe_heading}_diagram.svg"
            diagram_path = os.path.join(diagrams_dir, diagram_filename)

            # Generate high-quality PNG for document insertion with fallback
            png_path = diagram_path.replace('.svg', '.png')
            png_success = convert_mermaid_to_png_with_fallback(mermaid_code, png_path, heading_text)

            if png_success:
                # Use A4-optimized diagram insertion
                if insert_diagram_with_a4_optimization(doc, png_path, heading_text, diagram_type):
                    print(f"✅ Successfully processed diagram for: {heading_text}")
                    print(f"📄 A4-optimized PNG generated and inserted")

                    # Clean up PNG file after insertion
                    try:
                        os.remove(png_path)
                    except:
                        pass
                else:
                    print(f"⚠️ Failed to insert A4-optimized diagram for: {heading_text}")

                # Check if there are sub-diagrams to insert (only if main diagram was complex)
                sub_diagrams_for_section = [d for d in current_document_diagrams
                                          if d.get('parentSection') == heading_text and d.get('isSubDiagram')]

                print(f"🔍 DEBUG: Looking for sub-diagrams for '{heading_text}'")
                print(f"🔍 DEBUG: Found {len(sub_diagrams_for_section)} sub-diagrams")
                print(f"🔍 DEBUG: Total diagrams in memory: {len(current_document_diagrams)}")
                for i, diag in enumerate(current_document_diagrams):
                    print(f"   Diagram {i+1}: {diag.get('sectionTitle')} (parent: {diag.get('parentSection')}, isSubDiagram: {diag.get('isSubDiagram')})")

                if sub_diagrams_for_section:
                    print(f"📄 Inserting {len(sub_diagrams_for_section)} additional sub-diagrams...")
                    # Convert the stored sub-diagram data to the format expected by insert function
                    formatted_sub_diagrams = []
                    for sub_diag in sub_diagrams_for_section:
                        formatted_sub_diagrams.append({
                            'title': sub_diag['sectionTitle'],
                            'mermaid_code': sub_diag['mermaidCode'],
                            'focus_area': sub_diag.get('focusArea', 'System Component'),
                            'part_number': sub_diag.get('id', '').split('_')[-1] if '_' in sub_diag.get('id', '') else '1'
                        })
                    insert_sub_diagrams_into_document(doc, heading_text, formatted_sub_diagrams)
            else:
                print(f"⚠️ Failed to generate diagram for: {heading_text}")
                # Insert a placeholder instead of failing silently
                insert_diagram_placeholder(doc, heading_text, mermaid_code)
        else:
            print(f"⚠️ No Mermaid code generated for: {heading_text}")

def add_heading_content_with_full_context(doc: Document, heading: Dict[str, Any], uploaded_content: str = "", full_srs_content: str = ""):
    """Add content for a specific heading with diagram generation using full SRS context (fallback method)"""
    global current_document_diagrams

    heading_text = heading.get('heading', '')
    purpose = heading.get('purpose', '')
    source = heading.get('source', 'Unknown')
    category = heading.get('category', 'Other')
    user_prompt = heading.get('userPrompt', '')

    # Add heading (supports simple Markdown in titles)
    heading_para = doc.add_paragraph(style='SRS Heading 2')
    _add_markdown_runs(heading_para, heading_text)

    # Remove Purpose and Source metadata as requested
    # Add spacing after heading
    heading_para.paragraph_format.space_after = Pt(12)

    # Generate content using AI with uploaded document context and user_prompt
    print(f"📝 Generating content for heading: {heading_text}")
    generated_content = generate_content_for_heading(
        heading_text,
        purpose,
        source,
        uploaded_content,
        user_prompt
    )

    # Clean any remaining metadata from content
    generated_content = _clean_metadata_from_content(generated_content)

    # Add generated content with markdown formatting support
    _add_markdown_block(doc, generated_content, style_name='SRS Normal')

    # Check if this section should have a diagram (considering user prompt)
    diagram_type = should_generate_diagram(heading_text, user_prompt)
    if diagram_type:
        print(f"🎨 Generating diagram for section: {heading_text}")

        # Generate Mermaid diagram with custom prompt support and full SRS context
        mermaid_code = generate_mermaid_diagram(heading_text, generated_content, uploaded_content, user_prompt, full_srs_content)

        # If sequence diagram generation failed, try with a more specific prompt
        if not mermaid_code and diagram_type == 'sequenceDiagram':
            print(f"🔄 Sequence diagram generation failed, trying with specialized prompt...")
            mermaid_code = generate_sequence_diagram_specialized(heading_text, generated_content, uploaded_content, user_prompt, full_srs_content)

        if mermaid_code:
            # Convert Mermaid to PNG and insert into document
            # Create proper PNG file path
            safe_filename = heading_text.replace(' ', '_').replace('/', '_').replace('\\', '_')
            png_filename = f"{safe_filename}_diagram.png"
            png_path = os.path.join(os.path.dirname(__file__), '..', 'temp_diagrams', png_filename)

            # Ensure temp_diagrams directory exists
            os.makedirs(os.path.dirname(png_path), exist_ok=True)

            success = convert_mermaid_to_png(mermaid_code, png_path)
            if success and os.path.exists(png_path):
                # Use A4-optimized diagram insertion
                if insert_diagram_with_a4_optimization(doc, png_path, heading_text, diagram_type):
                    print(f"✅ Successfully processed diagram for: {heading_text}")
                    print(f"📄 A4-optimized PNG generated and inserted")

                    # Clean up PNG file after insertion
                    try:
                        os.remove(png_path)
                    except:
                        pass
                else:
                    print(f"⚠️ Failed to insert A4-optimized diagram for: {heading_text}")

                # Check if there are sub-diagrams to insert (only if main diagram was complex)
                sub_diagrams_for_section = [d for d in current_document_diagrams
                                          if d.get('parentSection') == heading_text and d.get('isSubDiagram')]

                if sub_diagrams_for_section:
                    print(f"📄 Inserting {len(sub_diagrams_for_section)} additional sub-diagrams...")
                    # Convert the stored sub-diagram data to the format expected by insert function
                    formatted_sub_diagrams = []
                    for sub_diag in sub_diagrams_for_section:
                        formatted_sub_diagrams.append({
                            'title': sub_diag['sectionTitle'],
                            'mermaid_code': sub_diag['mermaidCode'],
                            'focus_area': sub_diag.get('focusArea', 'System Component'),
                            'part_number': sub_diag.get('id', '').split('_')[-1] if '_' in sub_diag.get('id', '') else '1'
                        })
                    insert_sub_diagrams_into_document(doc, heading_text, formatted_sub_diagrams)
            else:
                print(f"⚠️ Failed to generate diagram for: {heading_text}")
                # Insert a placeholder instead of failing silently
                insert_diagram_placeholder(doc, heading_text, mermaid_code)
        else:
            print(f"⚠️ No Mermaid code generated for: {heading_text}")
    else:
        mermaid_code = None

        if mermaid_code:
            # Create temporary directory for diagrams if it doesn't exist
            diagrams_dir = os.path.join(os.path.dirname(__file__), '..', 'temp_diagrams')
            os.makedirs(diagrams_dir, exist_ok=True)

            # Generate unique filename for the diagram
            safe_heading = "".join(c for c in heading_text if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_heading = safe_heading.replace(' ', '_')
            diagram_filename = f"{safe_heading}_diagram.svg"
            diagram_path = os.path.join(diagrams_dir, diagram_filename)

            # Generate high-quality PNG for document insertion
            png_path = diagram_path.replace('.svg', '.png')
            png_success = convert_mermaid_to_png(mermaid_code, png_path)

            if png_success:
                # Add diagram description
                diagram_desc = doc.add_paragraph("The following diagram illustrates the system design:", style='SRS Normal')
                diagram_desc.paragraph_format.left_indent = Inches(0.25)
                diagram_desc.paragraph_format.space_after = Pt(6)

                # Insert high-quality PNG into document
                diagram_para = doc.add_paragraph()
                diagram_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = diagram_para.add_run()

                try:
                    run.add_picture(png_path, width=Inches(6))  # Full resolution for clarity
                    diagram_para.paragraph_format.space_after = Pt(12)
                    print(f"✅ Successfully inserted high-quality PNG diagram into document")

                    # Clean up PNG file after insertion
                    try:
                        os.remove(png_path)
                    except:
                        pass

                except Exception as png_error:
                    print(f"❌ Failed to insert PNG diagram: {png_error}")

                print(f"✅ Successfully processed diagram for: {heading_text}")
                print(f"📄 High-quality PNG generated and inserted")

                # Check if there are sub-diagrams to insert (only if main diagram was complex)
                sub_diagrams_for_section = [d for d in current_document_diagrams
                                          if d.get('parentSection') == heading_text and d.get('isSubDiagram')]

                if sub_diagrams_for_section:
                    print(f"📄 Inserting {len(sub_diagrams_for_section)} additional sub-diagrams...")
                    # Convert the stored sub-diagram data to the format expected by insert function
                    formatted_sub_diagrams = []
                    for sub_diag in sub_diagrams_for_section:
                        formatted_sub_diagrams.append({
                            'title': sub_diag['sectionTitle'],
                            'mermaid_code': sub_diag['mermaidCode'],
                            'focus_area': sub_diag.get('focusArea', 'System Component'),
                            'part_number': sub_diag.get('id', '').split('_')[-1] if '_' in sub_diag.get('id', '') else '1'
                        })
                    insert_sub_diagrams_into_document(doc, heading_text, formatted_sub_diagrams)
            else:
                print(f"⚠️ Failed to generate diagram for: {heading_text}")
        else:
            print(f"⚠️ No Mermaid code generated for: {heading_text}")

def generate_content_for_heading(heading: str, purpose: str, source: str, uploaded_content: str = "", user_prompt: str = "") -> str:
    """
    Generate content for a heading using Gemini AI, with optional user_prompt
    """
    try:
        print(f"📝 Generating content for heading: {heading}")

        # Configure Gemini API with error handling
        try:
            genai.configure(api_key="AIzaSyDS7AmlDi1cweQS1p-mJAUlB3uYHrYJXfI")
            model = genai.GenerativeModel('gemini-2.0-flash')
            print("✅ Gemini API configured for content generation")
        except Exception as config_error:
            print(f"❌ Failed to configure Gemini API: {config_error}")
            return f"Error: Failed to configure Gemini API for '{heading}'. Please check API key and internet connection."

        # Get heading-specific prompt
        heading_specific_prompt = get_heading_specific_prompt(heading)

        # Process user prompt to ensure professional SRS language
        if user_prompt and user_prompt.strip():
            # Clean up user prompt to be professional
            professional_prompt = user_prompt.replace("according to meeting summary", "based on system requirements")
            professional_prompt = professional_prompt.replace("meeting transcript", "system specifications")
            professional_prompt = professional_prompt.replace("meeting", "requirements analysis")
            professional_prompt = professional_prompt.replace("according to", "based on")

            prompt = f"""
            You are a professional SRS analyst writing to IEEE 29148 standard. Produce content that reads like it was authored by a human analyst, not AI.

            USER INPUT CONTEXT:
            {professional_prompt}

            SECTION HEADING: {heading}
            SYSTEM CONTEXT: {uploaded_content if uploaded_content else "No additional context available."}

            HEADING-SPECIFIC REQUIREMENTS:
            {heading_specific_prompt}

            {BASE_IEEE_DIRECTIVES}

7. Do not reference the source of requirements (e.g., “meeting notes,” “uploaded document”); write as if original to the SRS.
8. Reference diagrams textually only as "Figure x.y" without embedding them.
9. Preserve a professional tone and avoid casual expressions.
10. Not put spacing between two bullet points in the SRS, formatted to appear as a polished deliverable.
11. Ensure structure follows the expected Table of Contents for SRS (Introduction, Scope, Stakeholders, Requirements, Architecture, NFRs, etc.).
12. Expand each section with sufficient depth to resemble an analyst’s authored SRS, not a summary or checklist.
"""
        else:
            prompt = f"""
            You are a professional SRS analyst writing to IEEE 29148 standard. Produce content that reads like it was authored by a human analyst, not AI.

            SECTION HEADING: {heading}
            SYSTEM CONTEXT: {uploaded_content if uploaded_content else "No additional context available."}

            HEADING-SPECIFIC REQUIREMENTS:
            {heading_specific_prompt}

            {BASE_IEEE_DIRECTIVES}
1. Write in precise, formal prose suitable for contractual and technical documentation.
2. Use present tense and the keyword "shall" for all mandatory requirements.
3. Begin content immediately under the section heading with no prefatory labels or commentary.
4. Use structured paragraphs and bullet points (•) for requirements or subpoints; avoid numbered lists such as 1.1, 1.2, etc.
5. Ensure every requirement is atomic, testable, and unambiguous; avoid vague or compound statements.
6. Maintain consistency in terminology (system, actor, module, interface) across all sections.
7. Do not reference the source of requirements (e.g., “meeting notes,” “uploaded document”); write as if original to the SRS.
8. Reference diagrams textually only as "Figure x.y" without embedding them.
9. Preserve a professional tone and avoid casual expressions.
10. Not put spacing between two bullet points in the SRS, formatted to appear as a polished deliverable.
11. Ensure structure follows the expected Table of Contents for SRS (Introduction, Scope, Stakeholders, Requirements, Architecture, NFRs, etc.).
12. Expand each section with sufficient depth to resemble an analyst’s authored SRS, not a summary or checklist.
"""

        print(f"📤 Sending prompt to Gemini for '{heading}'...")

        # Add timeout and retry logic
        import time
        max_retries = 3
        timeout_seconds = 30

        for attempt in range(max_retries):
            try:
                print(f"🔄 Attempt {attempt + 1}/{max_retries} for '{heading}'...")

                # Set a timeout for the API call
                import signal
                def timeout_handler(signum, frame):
                    raise TimeoutError("Gemini API call timed out")

                # For Windows, we'll use a different approach since signal doesn't work well
                import threading
                result = [None]
                exception = [None]

                def api_call():
                    try:
                        result[0] = model.generate_content(prompt)
                    except Exception as e:
                        exception[0] = e

                thread = threading.Thread(target=api_call)
                thread.daemon = True
                thread.start()
                thread.join(timeout=timeout_seconds)

                if thread.is_alive():
                    print(f"⏰ Timeout after {timeout_seconds}s for '{heading}', retrying...")
                    continue

                if exception[0]:
                    raise exception[0]

                response = result[0]
                if response:
                    break

            except Exception as e:
                print(f"❌ Attempt {attempt + 1} failed for '{heading}': {str(e)}")
                if attempt == max_retries - 1:
                    print(f"💀 All attempts failed for '{heading}', using fallback content")
                    return f"Content generation failed for '{heading}' after {max_retries} attempts. Please try regenerating this section manually."

                # Wait before retry
                wait_time = (attempt + 1) * 2  # 2, 4, 6 seconds
                print(f"⏳ Waiting {wait_time}s before retry...")
                time.sleep(wait_time)

        if response and response.text:
            content = response.text.strip()
            print(f"Generated {len(content)} characters for '{heading}'")
            return content
        else:
            print(f"⚠️ Empty response from Gemini for '{heading}'")
            return f"No content generated for '{heading}'. Please try again."

    except Exception as e:
        print(f"❌ Error generating content for heading '{heading}': {e}")
        print(f"🔍 Error type: {type(e).__name__}")

        # Provide more specific error messages
        if "API_KEY_INVALID" in str(e):
            return f"Error: Invalid Gemini API key for '{heading}'. Please check your API key."
        elif "PERMISSION_DENIED" in str(e):
            return f"Error: Permission denied for Gemini API for '{heading}'. Check API key permissions."
        elif "QUOTA_EXCEEDED" in str(e):
            return f"Error: Gemini API quota exceeded for '{heading}'. Please try again later."
        else:
            return f"Error generating content for '{heading}': {str(e)}"



def add_page_number(paragraph):
    """Add page number to a paragraph"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def generate_diagram_signature(mermaid_code: str, heading: str) -> str:
    """
    Generate a unique signature for a diagram to prevent duplicates
    Returns a signature string that represents the diagram's content and structure
    """
    import re
    import hashlib

    if not mermaid_code:
        return ""

    # Extract key structural elements
    code_lower = mermaid_code.lower()

    # Get diagram type
    diagram_type = "unknown"
    if 'flowchart' in code_lower or 'graph' in code_lower:
        diagram_type = "flowchart"
    elif 'sequencediagram' in code_lower:
        diagram_type = "sequence"
    elif 'erdiagram' in code_lower:
        diagram_type = "er"
    elif 'classdiagram' in code_lower:
        diagram_type = "class"

    # Extract node labels (normalized)
    node_labels = re.findall(r'\[([^\]]+)\]', mermaid_code)
    normalized_labels = [re.sub(r'[^a-zA-Z0-9\s]', '', label.lower().strip()) for label in node_labels]
    normalized_labels.sort()

    # Extract connection patterns
    connections = re.findall(r'([A-Za-z0-9_]+)\s*-+>+\s*([A-Za-z0-9_]+)', mermaid_code)
    connection_patterns = [f"{a.lower()}-{b.lower()}" for a, b in connections]
    connection_patterns.sort()

    # Create signature components
    signature_parts = [
        f"type:{diagram_type}",
        f"heading:{heading.lower().replace(' ', '_')}",
        f"nodes:{len(normalized_labels)}",
        f"connections:{len(connection_patterns)}"
    ]

    # Add normalized node labels (limited to prevent signature explosion)
    if normalized_labels:
        node_signature = "_".join(normalized_labels[:5])  # Limit to first 5 nodes
        signature_parts.append(f"labels:{node_signature}")

    # Add connection patterns (limited)
    if connection_patterns:
        conn_signature = "_".join(connection_patterns[:5])  # Limit to first 5 connections
        signature_parts.append(f"flows:{conn_signature}")

    # Create final signature
    signature = "|".join(signature_parts)

    # Hash the signature for consistent length
    signature_hash = hashlib.md5(signature.encode()).hexdigest()[:16]

    return f"{diagram_type}_{signature_hash}"

def is_diagram_duplicate(mermaid_code: str, heading: str) -> bool:
    """
    Check if a diagram is a duplicate of an already generated one
    Returns True if duplicate, False if unique
    """
    global generated_diagram_signatures

    signature = generate_diagram_signature(mermaid_code, heading)
    print(f"🔍 Checking duplicate for '{heading}' - signature: {signature}")
    print(f"🔍 Current signatures: {list(generated_diagram_signatures)}")

    if signature in generated_diagram_signatures:
        print(f"⚠️ Duplicate diagram detected for '{heading}' (signature: {signature})")
        return True

    # Check for semantic similarity with existing diagrams
    for existing_sig in generated_diagram_signatures:
        if existing_sig.startswith(signature.split('_')[0]):  # Same diagram type
            # Check if nodes are too similar
            existing_nodes = set(existing_sig.split('|'))
            new_nodes = set(signature.split('|'))

            # Calculate similarity
            if existing_nodes and new_nodes:
                intersection = existing_nodes.intersection(new_nodes)
                union = existing_nodes.union(new_nodes)
                similarity = len(intersection) / len(union) if union else 0

                print(f"🔍 Similarity with {existing_sig}: {similarity:.1%}")

                if similarity > 0.8:  # Increased threshold to 80% to allow more variation
                    print(f"⚠️ Highly similar diagram detected for '{heading}' (similarity: {similarity:.1%})")
                    return True

    print(f"✅ No duplicate detected for '{heading}'")
    return False

def validate_diagram_for_a4(mermaid_code: str, heading: str) -> tuple[bool, str, dict]:
    """
    Validate that a diagram meets A4 compatibility requirements
    Returns: (is_valid: bool, error_message: str, complexity_analysis: dict)
    """
    if not mermaid_code or not mermaid_code.strip():
        return False, "Empty diagram code", {}

    # Analyze complexity
    complexity = analyze_diagram_complexity(mermaid_code)

    # Check A4 compatibility
    if complexity['is_too_complex']:
        error_msg = f"Diagram too complex for A4: {', '.join(complexity['suggestions'])}"
        return False, error_msg, complexity

    # Check for duplicate diagrams
    if is_diagram_duplicate(mermaid_code, heading):
        return False, "Duplicate diagram detected", complexity

    # Additional validation checks
    import re

    # Check for excessive text in nodes
    node_labels = re.findall(r'\[([^\]]+)\]', mermaid_code)
    long_labels = [label for label in node_labels if len(label) > 25]
    if long_labels:
        return False, f"Node labels too long for A4: {', '.join(long_labels[:3])}", complexity

    # Check for excessive connections
    connections = re.findall(r'-->|->>', mermaid_code)
    if len(connections) > 10:
        return False, f"Too many connections ({len(connections)}) for A4 page", complexity

    # Check for excessive subgraphs
    subgraphs = re.findall(r'subgraph\s+', mermaid_code)
    if len(subgraphs) > 3:
        return False, f"Too many subgraphs ({len(subgraphs)}) for A4 page", complexity

    return True, "Diagram is A4 compatible", complexity

def cleanup_diagram_for_a4(mermaid_code: str) -> str:
    """
    Clean up diagram code to ensure A4 compatibility
    """
    if not mermaid_code:
        return mermaid_code

    import re

    print(f"🔧 Cleaning diagram for A4 compatibility...")

    # Remove excessive whitespace and normalize
    lines = [line.strip() for line in mermaid_code.split('\n') if line.strip()]

    # Limit the number of nodes if too many
    node_pattern = r'([A-Za-z0-9_]+)\[([^\]]+)\]'
    nodes = re.findall(node_pattern, mermaid_code)

    if len(nodes) > 6:
        print(f"⚠️ Too many nodes ({len(nodes)}), limiting to 6 for A4 compatibility")
        # Keep only the first 6 nodes and their connections
        kept_nodes = set(node[0] for node in nodes[:6])

        # Filter lines to keep only those with kept nodes
        filtered_lines = []
        for line in lines:
            if any(node in line for node in kept_nodes):
                filtered_lines.append(line)
            elif line.startswith(('flowchart', 'graph', 'sequencediagram', 'erdiagram', 'classdiagram')):
                filtered_lines.append(line)
            elif line.startswith('subgraph') or line == 'end':
                filtered_lines.append(line)

        lines = filtered_lines

    # AGGRESSIVE CONNECTION LIMITING for A4 compatibility
    connection_pattern = r'([A-Za-z0-9_]+)\s*-+>+\s*([A-Za-z0-9_]+)'
    connections = re.findall(connection_pattern, mermaid_code)

    if len(connections) > 6:  # Limit to 6 connections maximum
        print(f"⚠️ Too many connections ({len(connections)}), limiting to 6 for A4 compatibility")
        # Keep only the first 6 connections
        kept_connections = connections[:6]

        # Filter lines to keep only those with kept connections
        filtered_lines = []
        for line in lines:
            # Keep diagram type and structural lines
            if line.startswith(('flowchart', 'graph', 'sequencediagram', 'erdiagram', 'classdiagram')):
                filtered_lines.append(line)
                continue
            elif line.startswith('subgraph') or line == 'end':
                filtered_lines.append(line)
                continue

            # Check if line contains a kept connection
            line_has_kept_connection = False
            for from_node, to_node in kept_connections:
                if from_node in line and to_node in line:
                    line_has_kept_connection = True
                    break

            if line_has_kept_connection:
                filtered_lines.append(line)
            elif '[' in line and ']' in line:  # Keep node definitions
                filtered_lines.append(line)

        lines = filtered_lines

    # Limit subgraphs if too many
    subgraph_count = sum(1 for line in lines if line.startswith('subgraph'))
    if subgraph_count > 2:
        print(f"⚠️ Too many subgraphs ({subgraph_count}), limiting to 2 for A4 compatibility")
        # Keep only the first 2 subgraphs
        subgraph_lines = []
        subgraph_count = 0
        in_subgraph = False

        for line in lines:
            if line.startswith('subgraph'):
                if subgraph_count < 2:
                    subgraph_lines.append(line)
                    subgraph_count += 1
                    in_subgraph = True
                else:
                    in_subgraph = False
            elif line == 'end' and in_subgraph:
                subgraph_lines.append(line)
                in_subgraph = False
            elif in_subgraph and subgraph_count <= 2:
                subgraph_lines.append(line)
            elif not in_subgraph:
                subgraph_lines.append(line)

        lines = subgraph_lines

    # Clean up node labels to be shorter but meaningful (NO ELLIPSIS)
    cleaned_lines = []
    for line in lines:
        # Shorten node labels if they're too long
        if '[' in line and ']' in line:
            # Find and shorten node labels properly without ellipsis
            def shorten_label(match):
                label = match.group(1)
                if len(label) > 15:
                    # Split into words and keep the most important ones
                    words = label.split()
                    if len(words) > 3:
                        # Keep first 2-3 most important words
                        shortened = ' '.join(words[:3])
                    elif len(words) > 1:
                        # Keep first 2 words
                        shortened = ' '.join(words[:2])
                    else:
                        # Single long word - keep first 12 characters
                        shortened = label[:12]
                    return f'[{shortened}]'
                return match.group(0)

            line = re.sub(r'\[([^\]]+)\]', shorten_label, line)

        cleaned_lines.append(line)

    result = '\n'.join(cleaned_lines)
    print(f"✅ Diagram cleaned for A4 compatibility")

    return result

def get_diagram_generation_summary() -> dict:
    """
    Get a summary of the current diagram generation status
    """
    global total_diagrams_generated, generated_diagram_signatures, current_document_diagrams

    summary = {
        'total_diagrams_generated': total_diagrams_generated,
        'max_diagrams_allowed': 3,
        'unique_signatures': len(generated_diagram_signatures),
        'current_diagrams': len(current_document_diagrams),
        'diagram_types': {},
        'a4_compatibility_status': 'Unknown'
    }

    # Analyze diagram types
    for diagram in current_document_diagrams:
        diagram_type = diagram.get('diagramType', 'unknown')
        if diagram_type not in summary['diagram_types']:
            summary['diagram_types'][diagram_type] = 0
        summary['diagram_types'][diagram_type] += 1

    # Check A4 compatibility
    if total_diagrams_generated == 0:
        summary['a4_compatibility_status'] = 'No diagrams generated'
    elif total_diagrams_generated >= 3:
        summary['a4_compatibility_status'] = 'Maximum diagrams reached'
    else:
        summary['a4_compatibility_status'] = 'Ready for more diagrams'

    return summary

def print_diagram_generation_summary():
    """
    Print a formatted summary of diagram generation status
    """
    summary = get_diagram_generation_summary()

    print("\n" + "="*60)
    print("📊 DIAGRAM GENERATION SUMMARY")
    print("="*60)
    print(f"📈 Total diagrams generated: {summary['total_diagrams_generated']}/{summary['max_diagrams_allowed']}")
    print(f"🔑 Unique diagram signatures: {summary['unique_signatures']}")
    print(f"📄 Current diagrams in memory: {summary['current_diagrams']}")
    print(f"✅ A4 compatibility status: {summary['a4_compatibility_status']}")

    if summary['diagram_types']:
        print(f"\n🎨 Diagram types generated:")
        for diagram_type, count in summary['diagram_types'].items():
            print(f"   • {diagram_type}: {count}")

    print("="*60 + "\n")

def generate_sequence_diagram_specialized(heading: str, content: str, uploaded_content: str = "", user_prompt: str = "", full_srs_content: str = "") -> str:
    """
    Specialized function to generate sequence diagrams when main generation fails
    Uses a more focused approach for sequence diagrams
    """
    try:
        print(f"🎯 Using specialized sequence diagram generation for: {heading}")

        # Configure Gemini API
        genai.configure(api_key="AIzaSyDS7AmlDi1cweQS1p-mJAUlB3uYHrYJXfI")
        model = genai.GenerativeModel('gemini-2.0-flash')

        # Specialized prompt for sequence diagrams
        prompt = f"""
        Generate a SIMPLE sequence diagram for this system section.

        SECTION: {heading}
        CONTENT: {content[:600]}
        FULL SRS CONTEXT: {full_srs_content[:1200] if full_srs_content else content[:1200]}

        REQUIREMENTS:
        1. Create a sequence diagram with MAXIMUM 3-4 participants
        2. Show only the MOST IMPORTANT interactions (4-6 messages maximum)
        3. Use simple participant names like "User", "System", "Database"
        4. Keep messages short and clear
        5. Focus on the core workflow only
        6. MUST fit on A4 page

        SEQUENCE DIAGRAM RULES:
        1. Start with: sequenceDiagram
        2. Declare participants: participant User, participant System, participant Database
        3. Use simple messages: User->>System: Request, System-->>User: Response
        4. Maximum 4-6 message exchanges
        5. No complex logic or loops

        EXAMPLE FORMAT:
        sequenceDiagram
            participant User
            participant System
            participant Database
            User->>System: Login Request
            System->>Database: Validate User
            Database-->>System: User Data
            System-->>User: Login Success

        Generate ONLY the sequence diagram code:
        """

        response = model.generate_content(prompt)
        if response and response.text:
            mermaid_code = response.text.strip()

            # Clean up the response
            if mermaid_code.startswith('```'):
                lines = mermaid_code.split('\n')
                mermaid_code = '\n'.join(lines[1:-1]) if len(lines) > 2 else mermaid_code

            # Apply syntax fixes
            fixed_code = fix_mermaid_syntax(mermaid_code)

            # Validate for A4 compatibility
            is_valid, error_msg, complexity = validate_diagram_for_a4(fixed_code, heading)
            if is_valid:
                print(f"✅ Specialized sequence diagram generated successfully")
                return fixed_code
            else:
                print(f"⚠️ Specialized sequence diagram still not A4 compatible: {error_msg}")
                return None
        else:
            print(f"⚠️ No response from specialized sequence diagram generation")
            return None

    except Exception as e:
        print(f"❌ Error in specialized sequence diagram generation: {e}")
        return None

if __name__ == "__main__":
    # Test the generator
    test_headings = [
        {
            'heading': '1. Introduction',
            'purpose': 'Briefly describe the system and its main objectives',
            'category': 'Introduction',
            'source': 'Standard Template'
        },
        {
            'heading': '2. Functional Requirements',
            'purpose': 'Define the functional requirements of the system',
            'category': 'Requirements',
            'source': 'Standard Template'
        },
        {
            'heading': 'Custom Section',
            'purpose': 'This is a custom section added by the user',
            'category': 'Custom',
            'source': 'Custom Section'
        }
    ]

    generate_srs_docx(test_headings, 'test_srs.docx')